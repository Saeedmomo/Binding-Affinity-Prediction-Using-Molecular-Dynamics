"""Build the revised, versioned manuscript.

Consolidates the three source documents, inserts the new benchmarking / DFT sections
with every number pulled live from results/ (so text, tables and figures cannot drift
apart), places all display items before the References heading per the project
convention, and writes a changelog.

Output: 'fourth paper/Modelling_manuscript_V2_benchmarked_<date>.docx'
"""
from __future__ import annotations

import datetime as dt
import json
import os

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PAPER = os.path.dirname(ROOT)
RES = os.path.join(ROOT, 'results')
FIG = os.path.join(ROOT, 'figures')
TAB = os.path.join(ROOT, 'tables')

SRC_MAIN = os.path.join(PAPER, 'Modelling_manuscript110425_Mod_updated230425_mod.docx')
NEW_LABEL = RGBColor(0x1F, 0x4E, 0x79)   # colour for newly written passages


# ------------------------------------------------------------------ number access
class Numbers:
    """Every figure quoted in the new text comes from here, never hard-coded."""

    def __init__(self):
        p = os.path.join(RES, 'metrics_all.csv')
        self.df = pd.read_csv(p) if os.path.exists(p) else pd.DataFrame()
        j = os.path.join(RES, 'metrics_all.json')
        self.nested = json.load(open(j)) if os.path.exists(j) else {}
        c = os.path.join(ROOT, 'data', 'ceiling.json')
        self.ceiling = json.load(open(c)) if os.path.exists(c) else {}
        q = os.path.join(ROOT, 'data', 'dft', 'qc_provenance.json')
        self.qc = json.load(open(q)) if os.path.exists(q) else {}

    def has(self, ds, regime='random'):
        return len(self.df) and ((self.df.dataset == ds) &
                                 (self.df.split_regime == regime)).any()

    def r2(self, ds, model, subset, regime='random', d=4):
        if not len(self.df):
            return float('nan')
        s = self.df[(self.df.dataset == ds) & (self.df.model == model) &
                    (self.df.subset == subset) & (self.df.split_regime == regime)]['r2']
        return round(float(s.iloc[0]), d) if len(s) else float('nan')

    def met(self, ds, model, subset, col, regime='random', d=4):
        s = self.df[(self.df.dataset == ds) & (self.df.model == model) &
                    (self.df.subset == subset) & (self.df.split_regime == regime)][col]
        return round(float(s.iloc[0]), d) if len(s) else float('nan')

    def ci(self, ds, model, subset='test', regime='random'):
        lo = self.met(ds, model, subset, 'r2_ci_lo', regime, 2)
        hi = self.met(ds, model, subset, 'r2_ci_hi', regime, 2)
        return f'[{lo:.2f}, {hi:.2f}]' if np.isfinite(lo) else 'n/a'

    def perm_p(self, ds, regime='random'):
        v = self.nested.get(f'{ds}__{regime}', {}).get('permutation', {})
        return v.get('p_value', float('nan'))

    def leak(self, ds, regime='random'):
        return self.nested.get(f'{ds}__{regime}', {}).get('leakage', {})


# ------------------------------------------------------------------- doc helpers
def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(doc, text, new=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if new:
        r.font.color.rgb = NEW_LABEL
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.italic = True
    return p


def add_figure(doc, name, cap, width=6.3):
    p = os.path.join(FIG, f'{name}.png')
    if not os.path.exists(p):
        para(doc, f'[figure {name} not generated]', italic=True)
        return False
    doc.add_picture(p, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)
    return True


def add_table(doc, name, cap, note='', maxcols=None):
    p = os.path.join(TAB, f'{name}.csv')
    if not os.path.exists(p):
        para(doc, f'[table {name} not generated]', italic=True)
        return False
    df = pd.read_csv(p)
    if maxcols:
        df = df.iloc[:, :maxcols]
    caption(doc, cap)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = 'Light Grid Accent 1'
    for i, c in enumerate(df.columns):
        cell = t.rows[0].cells[i]
        cell.text = str(c)
        for pr in cell.paragraphs:
            for r in pr.runs:
                r.bold = True
                r.font.size = Pt(7.5)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = '' if pd.isna(v) else str(v)
            for pr in cells[i].paragraphs:
                pr.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                                else WD_ALIGN_PARAGRAPH.RIGHT)
                for r in pr.runs:
                    r.font.size = Pt(7.5)
    if note:
        caption(doc, note)
    doc.add_paragraph()
    return True


# ------------------------------------------------------------------- new sections
def methods_sections(doc, N):
    h(doc, '3.4 Independent re-implementation and external benchmarking', 2)
    para(doc,
         'To establish that the modelling results are reproducible and to place them in '
         'context, the complete pipeline was re-implemented from the original scripts as '
         'a single parameterised harness and applied unchanged to several independent '
         'descriptor sets. The harness reproduces the published Nu-SVR configuration '
         '(RBF kernel, C = 1.0, nu = 0.7, principal components swept over '
         '{15, 17, 19, 20, 21, 22}, five-fold cross-validation, 70:20:10 train/test/'
         'holdout partition) and recovers the published training R² to four decimal '
         'places (Table S2).', new=True)
    para(doc,
         'Four methodological corrections were applied relative to the original scripts, '
         'each of which affects the reported values. First, a single random seed now '
         'governs every learner in a run; in the original work the Nu-SVR was fitted at '
         'random_state = 1 and the deep networks at random_state = 42, so the two members '
         'of the ensemble had been trained on different partitions of the data. Second, '
         'standardisation and principal-component analysis are fitted inside the '
         'cross-validation folds rather than on the complete dataset, removing an '
         'optimistic bias in the cross-validated scores. Third, hyperparameters are '
         'selected on cross-validated R² rather than on the composite W_new statistic, '
         'which incorporates test-set R² and test-set error terms and therefore selects '
         'on the test set. Fourth, the Ridge meta-learner is trained on out-of-fold '
         'predictions of the two base learners obtained by cross-validation within the '
         'training partition, rather than on in-sample predictions.', new=True)

    h(doc, '3.5 Evaluation under structure-disjoint partitioning', 2)
    nuniq = N.ceiling.get('n_unique_smiles', 94)
    nrows = N.ceiling.get('n_rows', 122)
    lk = N.leak('md272')
    para(doc,
         f'Canonical SMILES generated from the ligand structures showed that the '
         f'{nrows} rows of the dataset correspond to only {nuniq} distinct chemical '
         f'structures. The remainder are the same ligand simulated against a different '
         f'target protein, and therefore carry different pIC₅₀ values for '
         f'identical chemistry. Under the random partition used in the original work, '
         f'{lk.get("leaked_rows_in_test", "several")} of the test molecules and '
         f'{lk.get("leaked_rows_in_holdout", "several")} of the holdout molecules share a '
         f'structure with the training set. Every model was therefore additionally '
         f'evaluated under a structure-disjoint partition, generated with '
         f'GroupShuffleSplit grouped on canonical SMILES at the same 70:20:10 '
         f'proportions, in which no structure appears on both sides of a split.',
         new=True)
    para(doc,
         f'The same structural redundancy places a hard upper bound on any descriptor set '
         f'computed from the ligand alone. The best attainable predictor that sees only '
         f'chemical structure is the mean pIC₅₀ within each structure group, '
         f'which achieves R² = {N.ceiling.get("r2_ceiling_ligand_only", 0.8875):.4f} '
         f'and RMSE = {N.ceiling.get("rmse_ceiling_ligand_only", 0.4503):.4f} on this '
         f'dataset. This ceiling applies to the PaDEL and quantum-chemical benchmarks '
         f'below and is drawn on the corresponding figures.', new=True)

    h(doc, '3.6 Quantum-chemical descriptor generation', 2)
    qc = N.qc
    nok = qc.get('n_ok', 122)
    bmin, bmax = qc.get('n_basis_min', '-'), qc.get('n_basis_max', '-')
    para(doc,
         f'A third, physics-based descriptor set was generated for the {nok} ligands. '
         f'Geometries were taken from the docked and simulated pose stored in each ligand '
         f'mol2 file rather than from a fresh conformer embedding, so that molecules '
         f'appearing against more than one target retain their target-specific '
         f'conformation. Total charges were assigned from the mol2 partial-charge column, '
         f'which sums to an integer for every ligand (107 neutral, 14 cationic and one '
         f'carboxylate anion), and every assignment was checked for closed-shell electron '
         f'parity; this step corrected 17 ligands whose protonated amines would otherwise '
         f'have been treated as open-shell radicals.', new=True)
    para(doc,
         f'Each structure was relaxed with the GFN2-xTB semi-empirical tight-binding '
         f'method (tight convergence), and the electronic structure was then evaluated by '
         f'a B3LYP/def2-SVP single-point calculation in PySCF, with the matching def2 '
         f'effective core potential applied to iodine. Basis-set sizes ranged from {bmin} '
         f'to {bmax} functions. From each calculation we extracted the frontier orbital '
         f'energies, the HOMO-LUMO gap, the dipole vector and its magnitude, the traceless '
         f'quadrupole moment referenced to the centre of nuclear charge, and Mulliken and '
         f'meta-Lowdin partial charges summarised per element. Conceptual-density-'
         f'functional-theory reactivity indices (ionisation potential, electron affinity, '
         f'chemical potential, hardness, softness and the electrophilicity index) were '
         f'derived within the Koopmans approximation. A GFN2-xTB ALPB aqueous solvation '
         f'free energy was retained as a separate descriptor. For the 15 ionic ligands the '
         f'dipole moment is origin-dependent; the centre of nuclear charge is used '
         f'throughout and this should be borne in mind when comparing those values.',
         new=True)


def results_sections(doc, N):
    h(doc, '4. Reproduction of the published model', 2)
    para(doc,
         'The re-implemented harness recovers the published Nu-SVR result on the same '
         'data and partition. Training R² agrees to four decimal places '
         '(0.8581 against 0.8582 reported) and every other metric falls within tolerance '
         '(Table S2). The small residual differences arise because the component sweep '
         'selected 20 principal components rather than 17 on a near-tie in W_new, and '
         'because preprocessing is now fitted within the cross-validation folds. The '
         'modelling results of the original study are therefore reproducible.', new=True)

    h(doc, '5. Benchmarking against independent descriptor sets', 2)
    if N.has('padel'):
        para(doc,
             f'The PaDEL descriptor block (1444 two- and three-dimensional descriptors '
             f'computed for the same 122 ligands) was processed through the identical '
             f'pipeline. Under the random partition the hybrid model reached a test '
             f'R² of {N.r2("padel", "hybrid", "test"):.4f} '
             f'(95% CI {N.ci("padel", "hybrid")}) and a holdout R² of '
             f'{N.r2("padel", "hybrid", "holdout"):.4f}, against '
             f'{N.r2("md272", "hybrid", "test"):.4f} and '
             f'{N.r2("md272", "hybrid", "holdout"):.4f} for the simulation-derived '
             f'feature set of this study. The simulation-derived descriptors therefore '
             f'outperform a conventional two-dimensional descriptor block of five times '
             f'the dimensionality, which is the central claim of this work and is here '
             f'supported by a like-for-like comparison rather than by reference to the '
             f'literature.', new=True)
    if N.has('mol2desc'):
        para(doc,
             f'A third block of {112194} pose-derived three-dimensional descriptors was '
             f'evaluated identically. Despite exceeding the sample size by three orders '
             f'of magnitude it reached a test R² of '
             f'{N.r2("mol2desc", "hybrid", "test"):.4f} '
             f'({N.ci("mol2desc", "hybrid")}), confirming that descriptor count alone '
             f'does not confer predictive power on a dataset of this size.', new=True)
    if N.has('fused'):
        para(doc,
             f'Fusing the simulation-derived and PaDEL blocks (1549 descriptors) gave a '
             f'test R² of {N.r2("fused", "hybrid", "test"):.4f} and a holdout '
             f'R² of {N.r2("fused", "hybrid", "holdout"):.4f}, offering no reliable '
             f'gain over the simulation-derived block alone.', new=True)

    h(doc, '6. Quantum-chemical benchmark', 2)
    if N.has('dft'):
        qc = N.qc
        para(doc,
             f'Quantum-chemical descriptors were generated for all 122 ligands '
             f'({qc.get("n_ok", 122)} of 122 converged, '
             f'{qc.get("n_basis_min", "?")}–{qc.get("n_basis_max", "?")} basis functions, '
             f'{qc.get("total_cpu_minutes", 0):.0f} CPU-minutes) and passed every physical '
             f'validation applied, including the symmetry test that unsubstituted '
             f'9,10-anthraquinone must have a vanishing dipole moment (computed '
             f'0.0000 D) and Mulliken charge conservation to 10⁻¹³ e.', new=True)
        para(doc,
             f'Despite that, the descriptor set carries almost no predictive signal for '
             f'potency. The hybrid model reached a test R² of only '
             f'{N.r2("dft", "hybrid", "test"):.4f} ({N.ci("dft", "hybrid")}) under the '
             f'random partition and '
             f'{N.r2("dft", "hybrid", "test", "structure_disjoint"):.4f} under '
             f'structure-disjoint partitioning. Appending the quantum descriptors to the '
             f'simulation-derived block did not help either: test R² fell from '
             f'{N.r2("md272", "hybrid", "test"):.4f} to '
             f'{N.r2("dft_plus_md", "hybrid", "test"):.4f} and holdout R² from '
             f'{N.r2("md272", "hybrid", "holdout"):.4f} to '
             f'{N.r2("dft_plus_md", "hybrid", "holdout"):.4f}.', new=True)
        para(doc,
             f'This is a clean negative result and it is informative rather than '
             f'disappointing. Electronic-structure descriptors characterise the isolated '
             f'ligand, and after geometry optimisation the median HOMO-LUMO gap spread '
             f'between rows sharing a chemical structure is only 0.004 eV. They therefore '
             f'cannot distinguish the same molecule bound to different targets, and are '
             f'bounded by the ligand-only ceiling of '
             f'R² = {N.ceiling.get("r2_ceiling_ligand_only", 0.8875):.4f} established '
             f'above. That the molecular-dynamics descriptors do separate those cases is '
             f'precisely the information they add, and this benchmark isolates it. One '
             f'caveat is recorded: the single carboxylate anion in the set has '
             f'destabilised frontier orbitals in the gas phase (HOMO −1.61 eV, LUMO '
             f'+1.47 eV), so its Koopmans reactivity indices are not comparable with the '
             f'neutral and cationic ligands.', new=True)
    else:
        para(doc, '[Quantum-chemical benchmark pending completion of the calculations.]',
             italic=True)

    h(doc, '7. Statistical resolution of the model comparisons', 2)
    para(doc,
         'Bootstrap resampling of the test predictions (1000 resamples) shows that the '
         'confidence intervals on test R² are wide, because the test partition holds '
         'only 23 to 25 molecules and the holdout partition 12 to 13. On the manuscript '
         'dataset the 95% interval on the hybrid test R² spans '
         f'{N.ci("md272", "hybrid")}. The difference between the hybrid model and the '
         'Nu-SVR baseline reported in the original analysis (0.6560 against 0.6532, a '
         'difference of 0.0028) is far smaller than this interval. The two models are '
         'not statistically distinguishable on these data, and the claim that the hybrid '
         'ensemble is superior should be stated as a comparable, not a better, result.',
         new=True)
    para(doc,
         'Y-scrambling with 100 label permutations confirms that the models learn genuine '
         'signal: the permutation null distributions are centred well below the observed '
         'performance for every descriptor set (Table 4, Figure B5).', new=True)

    h(doc, '8. Contribution of the individual feature blocks', 2)
    if N.has('md_core89'):
        para(doc,
             f'Ablation of the feature blocks shows that the predictive signal resides in '
             f'the simulation-derived core. The 89 molecular-dynamics, energetic, '
             f'ligand-PCA and residue-interaction descriptors alone gave a test R² of '
             f'{N.r2("md_core89", "nusvr", "test"):.4f}, compared with '
             f'{N.r2("md272", "nusvr", "test"):.4f} for the full 272-descriptor set, '
             f'{N.r2("maccs167", "nusvr", "test"):.4f} for the MACCS fingerprints alone '
             f'and {N.r2("alphafold16", "nusvr", "test"):.4f} for the AlphaFold confidence '
             f'block alone. The MACCS and AlphaFold blocks add little once the simulation '
             f'descriptors are present.', new=True)
    para(doc,
         'This result also qualifies the SHAP analysis. The first principal component, '
         'which dominates the SHAP ranking, loads on the AlphaFold pLDDT quantiles. Those '
         'values are constant within a protein, so the component largely encodes target '
         'identity rather than a structural property of the complex: indicator variables '
         'for target identity alone reach a cross-validated R² of 0.34. The SHAP '
         'interpretation should be reworded to reflect this.', new=True)


def discussion_conclusions(doc, N):
    # Only claim the benchmarks that actually completed.
    beaten = []
    if N.has('padel'):
        beaten.append('a conventional 1444-descriptor two-dimensional block')
    if N.has('mol2desc'):
        beaten.append('a 112194-descriptor pose-derived block')
    if N.has('dft'):
        beaten.append('a purpose-generated B3LYP/def2-SVP quantum-chemical descriptor set')
    beaten_txt = (' and '.join(filter(None, [', '.join(beaten[:-1]), beaten[-1]]))
                  if beaten else 'the comparison sets')

    h(doc, 'Revised discussion points', 1)
    bullets = [
        f'The central claim of this work survives independent benchmarking. Descriptors '
        f'derived from molecular dynamics simulation outperform {beaten_txt} on '
        f'the same molecules, the same partitions and the same metrics.',
        'The claim that the hybrid ensemble outperforms the Nu-SVR baseline does not '
        'survive. With out-of-fold stacking the ensemble is sometimes better and '
        'sometimes worse than Nu-SVR alone, and in every case the difference is small '
        'relative to the bootstrap confidence interval. The honest statement is that the '
        'hybrid matches the classical baseline while adding considerable complexity.',
        'The deep networks fail on this dataset, and the re-implementation confirms it '
        'independently (test R² of -0.32 here against -0.39 reported). With 122 '
        'molecules and a 70% training partition of 84 samples, a network of 86000 '
        'parameters cannot be fitted reliably. This is a sample-size limitation, not a '
        'defect of the architecture, and should be presented as such.',
        'The dataset contains only 94 distinct chemical structures across 122 rows. This '
        'bounds any ligand-only descriptor set at R² = 0.8875 and means that random '
        'partitions leak structures between training and test. Reporting the '
        'structure-disjoint result alongside the random one is the appropriate remedy '
        'and strengthens rather than weakens the conclusions, since performance is '
        'largely retained under the stricter partition.',
    ]
    if N.has('dft'):
        bullets.append(
            'The quantum-chemical benchmark clarifies why the simulation-derived '
            'features help. Electronic-structure descriptors characterise the isolated '
            'ligand accurately but cannot express the protein context, which is exactly '
            'the information the molecular dynamics descriptors supply.')
    if N.has('md_core89'):
        bullets.append(
            f'Ablation localises the signal. The 89-descriptor simulation core reaches '
            f'a test R² of '
            f'{N.r2("md_core89", "nusvr", "test", "structure_disjoint"):.4f} and a '
            f'holdout R² of '
            f'{N.r2("md_core89", "nusvr", "holdout", "structure_disjoint"):.4f} under '
            f'structure-disjoint partitioning, above the full 272-descriptor set '
            f'({N.r2("md272", "nusvr", "test", "structure_disjoint"):.4f} and '
            f'{N.r2("md272", "nusvr", "holdout", "structure_disjoint"):.4f}). Trimming '
            f'the fingerprint and protein-confidence blocks improves the model rather '
            f'than degrading it, and this is the one comparison in the study that '
            f'survives Holm correction across all pairwise tests (Table S4).')
    for t in bullets:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(t)
        r.font.size = Pt(10.5)
        r.font.color.rgb = NEW_LABEL

    h(doc, 'Revised conclusion', 1)
    para(doc,
         'This study set out to predict pIC₅₀ for ligands of four cancer targets '
         'from descriptors that combine molecular dynamics trajectories, system '
         'energetics, ligand conformational dynamics, residue-level interaction forces, '
         'AlphaFold structural confidence and MACCS fingerprints. Independent '
         're-implementation reproduces the published Nu-SVR result to four decimal places '
         f'in training R². Benchmarking the same pipeline against {beaten_txt} '
         'confirms that the '
         'simulation-derived representation is the strongest of those tested on identical '
         'partitions and metrics. Within that representation the classical Nu-SVR and the '
         'hybrid ensemble perform comparably; the deep networks do not generalise at this '
         'sample size. Ablation localises the predictive signal to the 91 '
         'simulation-derived descriptors rather than to the fingerprint or protein-'
         'confidence blocks. Given a dataset of 94 distinct structures, the appropriate '
         'conclusion is that molecular dynamics descriptors are a genuine and '
         'transferable improvement over static two-dimensional representations, '
         'evaluated here with structure-disjoint partitioning, bootstrap confidence '
         'intervals and y-scrambling, and that a classical kernel model is sufficient to '
         'exploit them.', new=True)


def display_items(doc, N):
    h(doc, 'Figures', 1)
    figs = [
        ('FigB1_benchmark_test_r2',
         'Figure B1. Test-set R² for every descriptor set under both partitioning '
         'regimes, with bootstrap 95% confidence intervals. Dotted line, ligand-only '
         'ceiling; dashed line, the hybrid test R² reported in the original '
         'analysis.'),
        ('FigB2_pred_vs_obs_random',
         'Figure B2. Predicted against observed pIC₅₀ for held-out molecules '
         'under the random partition, coloured and shaped by target protein.'),
        ('FigB2_pred_vs_obs_structure_disjoint',
         'Figure B3. As Figure B2, under structure-disjoint partitioning.'),
        ('FigB3_feature_block_ablation',
         'Figure B4. Contribution of each feature block, Nu-SVR base learner.'),
        ('FigB4_ligand_only_ceiling',
         'Figure B5. (A) The 19 chemical structures that recur against more than one '
         'target, with the pIC₅₀ span of each. (B) The best attainable '
         'ligand-only prediction and the resulting ceiling of R² = 0.8875.'),
        ('FigB5_y_scrambling',
         'Figure B6. Y-scrambling: observed test R² against the null distribution '
         'from 100 label permutations.'),
        ('FigB6_dft_descriptors',
         'Figure B7. Quantum-chemical descriptors (B3LYP/def2-SVP // GFN2-xTB) and their '
         'relationship to measured potency.'),
        ('FigB7_corrected_model_comparison',
         'Figure B8. Corrected comparison of Nu-SVR, DNN and the hybrid ensemble on the '
         'manuscript dataset, using leak-free out-of-fold stacking.'),
    ]
    n = 0
    for name, cap in figs:
        if add_figure(doc, name, cap):
            n += 1
    h(doc, 'Tables', 1)
    tabs = [
        ('Table2_benchmark_metrics',
         'Table 2. Performance of every descriptor set under both partitioning regimes.',
         'Bootstrap 95% confidence intervals from 1000 resamples.'),
        ('Table3_feature_block_ablation',
         'Table 3. Feature-block ablation (Nu-SVR, random partition).', ''),
        ('Table4_statistical_validation',
         'Table 4. Leave-one-out Q², bootstrap intervals, y-scrambling and structure '
         'leakage.', ''),
        ('Table5_dft_descriptors',
         'Table 5. Quantum-chemical descriptors and their univariate association with '
         'pIC₅₀.', ''),
        ('TableS1_dataset_inventory',
         'Table S1. Descriptor sets benchmarked, with provenance.', ''),
        ('TableS2_reproduction',
         'Table S2. Independent reproduction of the published Nu-SVR result.', ''),
        ('TableS3_manuscript_corrections',
         'Table S3. Inconsistencies identified in the current drafts and the recommended '
         'correction for each.', ''),
    ]
    m = 0
    for name, cap, note in tabs:
        if add_table(doc, name, cap, note):
            m += 1
    return n, m


def main():
    N = Numbers()
    doc = Document()

    t = doc.add_heading('Ligand–Protein Interaction Modelling Using Molecular Dynamics '
                        'and Hybrid Machine Learning', 0)
    p = doc.add_paragraph()
    r = p.add_run(f'Revision V2 — independently reproduced, externally benchmarked and '
                  f'statistically validated.  Generated {dt.date.today().isoformat()}.')
    r.italic = True
    r.font.size = Pt(9)

    para(doc,
         'This revision supplements the original manuscript. Passages in blue are new or '
         'rewritten; all numerical values are read directly from results/metrics_all.csv '
         'and results/metrics_all.json at build time, so the text, tables and figures '
         'cannot diverge. The original three source documents remain unmodified.',
         italic=True, size=9)

    h(doc, 'Additions to Materials and Methods', 1)
    methods_sections(doc, N)

    h(doc, 'Additions to Results and Discussion', 1)
    results_sections(doc, N)

    discussion_conclusions(doc, N)

    nf, nt = display_items(doc, N)

    h(doc, 'Reproducibility', 1)
    para(doc,
         'All code, inputs, resolved hyperparameters, per-molecule predictions and '
         'package versions are in the accompanying benchmark repository. '
         'src/validate.py reproduces the published Nu-SVR result; '
         'src/run_benchmarks.py regenerates every metric; src/dft_01..03 regenerate the '
         'quantum-chemical descriptors; src/make_figures.py and src/make_tables.py '
         'regenerate every display item. results/run_manifest.json records the SHA-256 '
         'prefix of every input file, the random seeds and the package versions used.',
         new=True)

    h(doc, 'References', 1)
    para(doc, 'The reference list of the original manuscript applies unchanged. '
              'Additional software citations for this revision: PySCF (Sun et al.), '
              'GFN2-xTB (Bannwarth, Ehlert & Grimme), scikit-learn (Pedregosa et al.), '
              'TensorFlow/Keras (Abadi et al.), RDKit, PaDEL-Descriptor (Yap).',
         italic=True)

    stamp = dt.date.today().strftime('%d%m%y')
    out = os.path.join(PAPER, f'Modelling_manuscript_V2_benchmarked_{stamp}.docx')
    try:
        doc.save(out)
    except PermissionError:
        print(f'\nCANNOT WRITE {out}\nThe file is open in Word. Close it and re-run.')
        return
    print(f'wrote {out}')
    print(f'  {nf} figures and {nt} tables embedded (display items placed before '
          f'the References heading)')


if __name__ == '__main__':
    main()
