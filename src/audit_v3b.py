"""V3b audit. Fails loudly.

Extends audit_v3 with four checks that the V3 gate did not make:
  supporting items must appear in body order matching their numbers
  every numbered supporting item must be cited somewhere in the manuscript body
  no retired title or heading string may survive in either document
  the supporting information subtitle must equal the manuscript title
"""
import glob
import os
import re
import sys

from docx import Document

V3B = r'D:\Chua files\fourth paper\V3b'

TITLE = ('A Controlled Benchmark of Molecular Dynamics Derived Descriptors for '
         'Ligand Potency Prediction')

# strings that must not survive anywhere in either document
RETIRED = (
    'Molecular Dynamics Derived Descriptors Outperform',
    'Simulation-Derived Descriptors Outperform Every Comparison Set',
    'The Predictive Signal Resides in the Simulation-Derived Core',
    'Which Simulation-Derived Descriptors Carry the Information',
    'Model Choice Is Secondary to Representation',
    'The remaining supporting tables follow the order given in the Contents',
    'pose-derived', 'Pose 3D', 'mol2desc',
)


def read(path):
    d = Document(path)
    paras = [q.text for q in d.paragraphs]
    cells = [c.text for t in d.tables for r in t.rows for c in r.cells]
    return d, paras, cells


def main():
    files = sorted(glob.glob(os.path.join(V3B, '*.docx')))
    if len(files) != 2:
        print(f'FAIL expected 2 documents in V3b, found {len(files)}')
        return 1

    ok = True
    si_numbered = {}          # 'Table S3' -> caption text, filled from the SI
    manuscript_body = ''
    si_subtitle = None

    for path in files:
        name = os.path.basename(path)
        is_si = 'supporting' in name.lower()
        d, paras, cells = read(path)
        allt = paras + cells
        body = '\n'.join(allt)
        if not is_si:
            manuscript_body = body

        print(f'=== {name} ===')
        print(f'  paragraphs {len(paras)}  tables {len(d.tables)}  '
              f'images {len(d.inline_shapes)}  words ~{len(body.split())}')

        # ---------------------------------------------------- typography rules
        for ch, label in (('\u2014', 'em dash'), ('\u2013', 'en dash'),
                          ('_', 'underscore')):
            hits = [t for t in allt if ch in t]
            if hits:
                ok = False
                print(f'  FAIL {label}: {len(hits)}')
                for h in hits[:3]:
                    i = h.index(ch)
                    print(f'       ...{h[max(0, i - 50):i + 50]}...')

        # ------------------------------------------------------ retired strings
        for s in RETIRED:
            if s in body:
                ok = False
                print(f'  FAIL retired string present: {s!r}')

        # ------------------------------------------- unnumbered SI references
        for t in paras:
            s = t.strip()
            if not re.search(r'\b(Supporting|Supplementary) Information\b', s):
                continue
            if is_si and s.startswith('Supporting Information'):
                continue                                   # the SI's own title
            if 'available free of charge' in s:
                continue                                   # Associated Content
            if not re.search(r'\b(Table|Figure|Section)\s+S\d+', s):
                ok = False
                print(f'  FAIL unnumbered reference: {s[:150]}')

        if 'PyDescriptor' not in body:
            ok = False
            print('  FAIL PyDescriptor not named')

        if not is_si:
            if paras[0].strip() != TITLE:
                ok = False
                print(f'  FAIL title is {paras[0].strip()[:90]!r}')
            for need in ('Q(H)', 'E(ionic)', 'E(cation-pi)', 'dG = gamma A',
                         'W(new)'):
                if need not in body:
                    ok = False
                    print(f'  FAIL missing equation element: {need!r}')
            for cap in ('Table 1.', 'Table 2.', 'Table 3.'):
                if cap not in body:
                    ok = False
                    print(f'  FAIL missing main-text {cap}')
            if 'Table S3. Performance' in body:
                ok = False
                print('  FAIL the metrics table is still in the main text')
        else:
            si_subtitle = paras[1].strip() if len(paras) > 1 else ''

            # The Contents listing repeats every caption in the correct order at
            # the top of the document. Taking first occurrences therefore reads
            # the Contents rather than the body and passes unconditionally, which
            # is what the first version of this gate did. Skip everything up to
            # the heading that follows Contents, then read real captions only.
            start = 0
            for i, q in enumerate(d.paragraphs):
                if q.style.name.startswith('Heading') and q.text.strip() == 'Contents':
                    start = i + 1
                elif start and q.style.name.startswith('Heading'):
                    start = i
                    break
            body_paras = [q.text for q in d.paragraphs[start:]]
            if start == 0:
                ok = False
                print('  FAIL could not locate the Contents section')

            for cls in ('Table', 'Figure'):
                order = []
                for t in body_paras:
                    m = re.match(rf'{cls} S(\d+)\.', t.strip())
                    if m:
                        order.append(int(m.group(1)))
                        si_numbered[f'{cls} S{m.group(1)}'] = t.strip()
                dupes = sorted({x for x in order if order.count(x) > 1})
                if dupes:
                    ok = False
                    print(f'  FAIL duplicate {cls} S captions in body: {dupes}')
                if not order:
                    continue
                gaps = sorted(set(range(1, max(order) + 1)) - set(order))
                if gaps:
                    ok = False
                    print(f'  FAIL {cls} S numbering has gaps: {gaps}')
                elif order != sorted(order):
                    ok = False
                    print(f'  FAIL {cls} S captions out of body order: {order}')
                else:
                    print(f'  ok {cls}s S1 to S{max(order)} present and in order')
        print()

    # -------------------------------------------------- SI subtitle matches title
    if si_subtitle != TITLE:
        ok = False
        print(f'FAIL supporting information subtitle is {si_subtitle!r}')

    # ----------------------------------- every supporting item cited in the body
    # digit boundary, so that citing Table S1 does not satisfy Table S10
    def cited(key):
        cls, num = key.split(' S')
        return re.search(rf'\b{cls}s?\s+S{num}(?!\d)', manuscript_body) is not None

    uncited = [k for k in sorted(si_numbered) if not cited(k)]
    if uncited:
        ok = False
        print(f'FAIL supporting items never cited in the manuscript: '
              f'{", ".join(uncited)}')
    else:
        print(f'ok all {len(si_numbered)} supporting items cited in the manuscript')

    print('PASS' if ok else 'FAILED')
    return 0 if ok else 1


if __name__ == '__main__':
    sys.exit(main())
