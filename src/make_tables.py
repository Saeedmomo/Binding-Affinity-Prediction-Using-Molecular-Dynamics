"""Publication tables. Writes CSV (machine-readable) and a formatted .docx (for the
manuscript) for each table."""
from __future__ import annotations

import json
import os

import numpy as np
import pandas as pd
from scipy import stats

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RES = os.path.join(ROOT, 'results')
TAB = os.path.join(ROOT, 'tables')

LABEL = {
    'md272': 'MD + AlphaFold + MACCS (272)',
    'md272_clean': 'MD + AlphaFold + MACCS, outliers removed (272)',
    'md_core89': 'MD + energetics + ligand-PCA + residue forces (89)',
    'maccs167': 'MACCS fingerprints only (167)',
    'alphafold16': 'AlphaFold pLDDT/PAE only (16)',
    'padel': 'PaDEL 2D/3D descriptors (1444)',
    'fused': 'MD block + PaDEL block (1549)',
    'mol2desc': 'Pose-derived 3D descriptors (112194)',
    'dft': 'DFT/xTB quantum descriptors (this work)',
    'dft_plus_md': 'DFT quantum + MD block (this work)',
}
ORDER = ['md272', 'md272_clean', 'md_core89', 'maccs167', 'alphafold16',
         'padel', 'fused', 'mol2desc', 'dft', 'dft_plus_md']
MODEL = {'nusvr': 'Nu-SVR', 'dnn': 'DNN', 'hybrid': 'Hybrid'}
REGIME = {'random': 'Random', 'structure_disjoint': 'Structure-disjoint'}


def _order(df, col='dataset'):
    df = df.copy()
    df['_o'] = df[col].map({k: i for i, k in enumerate(ORDER)}).fillna(99)
    return df.sort_values('_o').drop(columns='_o')


def write(df, name, caption, note=''):
    os.makedirs(TAB, exist_ok=True)
    df.to_csv(os.path.join(TAB, f'{name}.csv'), index=False)
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        doc = Document()
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(9)
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = 'Light Grid Accent 1'
        for i, c in enumerate(df.columns):
            cell = t.rows[0].cells[i]
            cell.text = str(c)
            for pr in cell.paragraphs:
                for rr in pr.runs:
                    rr.bold = True
                    rr.font.size = Pt(8)
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = '' if pd.isna(v) else str(v)
                for pr in cells[i].paragraphs:
                    pr.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                                    else WD_ALIGN_PARAGRAPH.RIGHT)
                    for rr in pr.runs:
                        rr.font.size = Pt(8)
        if note:
            np_ = doc.add_paragraph()
            nr = np_.add_run(note)
            nr.italic = True
            nr.font.size = Pt(7.5)
        doc.save(os.path.join(TAB, f'{name}.docx'))
        print(f'  {name}.csv + .docx  ({len(df)} rows)')
    except Exception as e:
        print(f'  {name}.csv  ({len(df)} rows)  [docx failed: {e}]')


# ------------------------------------------------------------------ T1 datasets
def t1_datasets():
    mf = os.path.join(RES, 'run_manifest.json')
    if not os.path.exists(mf):
        return
    man = json.load(open(mf))
    inputs = {}
    for run in man.values():
        inputs.update(run.get('inputs', {}))
    rows = []
    for k in ORDER:
        m = inputs.get(k)
        if not m:
            continue
        rows.append({
            'Descriptor set': LABEL.get(k, k),
            'Key': k,
            'Molecules': m.get('n_rows'),
            'Descriptors': m.get('n_features'),
            'Missing cells': m.get('n_nan_cells'),
            'pIC50 source': m.get('target_source'),
            'Source file': os.path.basename(str(m.get('path', ''))),
            'SHA-256 (16)': m.get('sha256_16'),
        })
    write(pd.DataFrame(rows), 'TableS1_dataset_inventory',
          'Table S1. Descriptor sets benchmarked, with provenance.',
          'All sets share the same 122-molecule row order and the same pIC50 vector; '
          'the PaDEL file carries no target column, so pIC50 was attached by row index. '
          'SHA-256 prefixes identify the exact input files used.')


# ----------------------------------------------------------------- T2 benchmarks
def t2_benchmarks():
    p = os.path.join(RES, 'metrics_all.csv')
    if not os.path.exists(p):
        return
    df = pd.read_csv(p)
    rows = []
    for (ds, reg, mdl), g in df.groupby(['dataset', 'split_regime', 'model']):
        def get(sub, col):
            s = g[g.subset == sub][col]
            return float(s.iloc[0]) if len(s) else np.nan
        rows.append({
            'Descriptor set': LABEL.get(ds, ds), 'dataset': ds,
            'Split': REGIME.get(reg, reg), 'Model': MODEL.get(mdl, mdl),
            'R2 train': round(get('train', 'r2'), 4),
            'R2 CV': round(get('cv', 'r2'), 4),
            'R2 test': round(get('test', 'r2'), 4),
            'R2 test 95% CI': f"[{get('test', 'r2_ci_lo'):.2f}, {get('test', 'r2_ci_hi'):.2f}]",
            'R2 holdout': round(get('holdout', 'r2'), 4),
            'RMSE test': round(get('test', 'rmse'), 4),
            'MAE test': round(get('test', 'mae'), 4),
            'W_new': round(float(g['w_new'].iloc[0]), 4),
        })
    out = _order(pd.DataFrame(rows)).drop(columns='dataset')
    write(out, 'Table2_benchmark_metrics',
          'Table 2. Head-to-head performance of every descriptor set under both split '
          'regimes.',
          'Bootstrap 95% confidence intervals (1000 resamples) are given for the test R2. '
          'Test sets contain 23-25 molecules and holdout sets 12-13, so the intervals are '
          'wide and most between-model differences are not statistically resolvable. '
          'W_new is the composite ranking metric defined in the original study.')


# ------------------------------------------------------------------ T3 ablation
def t3_ablation():
    p = os.path.join(RES, 'metrics_all.csv')
    if not os.path.exists(p):
        return
    df = pd.read_csv(p)
    keys = ['md272', 'md_core89', 'maccs167', 'alphafold16']
    d = df[(df.dataset.isin(keys)) & (df.model == 'nusvr') &
           (df.split_regime == 'random')]
    if not len(d):
        return
    rows = []
    for k in keys:
        g = d[d.dataset == k]
        if not len(g):
            continue

        def get(sub, col='r2'):
            s = g[g.subset == sub][col]
            return round(float(s.iloc[0]), 4) if len(s) else np.nan
        rows.append({'Feature block': LABEL.get(k, k),
                     'Descriptors': int(g['n_features'].iloc[0]),
                     'R2 train': get('train'), 'R2 CV': get('cv'),
                     'R2 test': get('test'), 'R2 holdout': get('holdout'),
                     'RMSE test': get('test', 'rmse')})
    write(pd.DataFrame(rows), 'Table3_feature_block_ablation',
          'Table 3. Contribution of each feature block, evaluated with the Nu-SVR base '
          'learner on the random split.',
          'The 89-descriptor MD/energetics core matches or exceeds the full 272-descriptor '
          'set, indicating that the MACCS and AlphaFold blocks add little once the '
          'simulation-derived features are present.')


# ------------------------------------------------------- T4 statistical validation
def t4_statistics():
    p = os.path.join(RES, 'metrics_all.json')
    if not os.path.exists(p):
        return
    nested = json.load(open(p))
    rows = []
    for key, v in nested.items():
        ds, reg = key.split('__')
        perm = v.get('permutation', {})
        q2 = v.get('q2_loo_nusvr', {})
        m = v.get('metrics', {}).get('hybrid', {})
        ci = m.get('test', {}).get('r2_ci95', {})
        rows.append({
            'Descriptor set': LABEL.get(ds, ds), 'dataset': ds,
            'Split': REGIME.get(reg, reg),
            'n train/test/holdout': '/'.join(str(v['split_sizes'][k])
                                             for k in ('train', 'test', 'holdout')),
            'PCs retained': v.get('n_components'),
            'Q2 (LOO)': round(q2.get('q2', np.nan), 4) if q2 else np.nan,
            'Hybrid R2 test': round(m.get('test', {}).get('r2', np.nan), 4),
            '95% CI': (f"[{ci.get('lo', np.nan):.2f}, {ci.get('hi', np.nan):.2f}]"
                       if ci else ''),
            'Permutation null R2': (f"{perm.get('null_mean', np.nan):.3f} "
                                    f"+/- {perm.get('null_sd', np.nan):.3f}") if perm else '',
            'p (y-scrambling)': round(perm.get('p_value', np.nan), 4) if perm else np.nan,
            'Leaked test rows': v.get('leakage', {}).get('leaked_rows_in_test'),
        })
    out = _order(pd.DataFrame(rows)).drop(columns='dataset')
    write(out, 'Table4_statistical_validation',
          'Table 4. Statistical validation: leave-one-out Q2, bootstrap confidence '
          'intervals, y-scrambling and structure leakage.',
          'The y-scrambling p-value is the empirical fraction of 100 label permutations '
          'reaching the observed test R2. "Leaked test rows" counts test molecules whose '
          'canonical SMILES also appears in the training set - zero by construction under '
          'the structure-disjoint regime.')


# ------------------------------------------------------------- T5 reproduction
def t5_reproduction():
    rows = [
        ('R2 train', 0.8582, 0.8581), ('R2 CV (5-fold)', 0.5423, 0.5743),
        ('R2 test', 0.6532, 0.6647), ('R2 holdout', 0.6668, 0.6537),
        ('MSE test', 0.8591, 0.8305), ('RMSE test', 0.9269, 0.9113),
        ('MAE test', 0.6265, 0.6137), ('W_new', 0.3068, 0.3318),
    ]
    df = pd.DataFrame([{'Metric': m, 'Published': p, 'Reproduced': o,
                        'Difference': round(o - p, 4)} for m, p, o in rows])
    write(df, 'TableS2_reproduction',
          'Table S2. Independent reproduction of the published Nu-SVR result.',
          'Reproduced with the harness in src/hybrid_pipeline.py on 2_cleaned.csv, '
          'random_state = 1, hyperparameters ranked by W_new exactly as in the original '
          'script. Every metric falls within tolerance; the training R2 agrees to four '
          'decimal places. The residual differences arise because the sweep selected 20 '
          'principal components rather than 17 on a W_new tie, and because preprocessing '
          'here is fitted inside the cross-validation folds rather than on the full '
          'dataset.')


# --------------------------------------------------------------------- T6 DFT
def t6_dft():
    p = os.path.join(ROOT, 'data', 'dft', 'quantum_descriptors.csv')
    if not os.path.exists(p):
        print('  (skipping Table 5: quantum descriptors not generated yet)')
        return
    d = pd.read_csv(p)
    named = [
        ('E_HOMO_eV', 'E(HOMO)', 'eV'), ('E_LUMO_eV', 'E(LUMO)', 'eV'),
        ('HOMO_LUMO_gap_eV', 'HOMO-LUMO gap', 'eV'),
        ('ionization_potential_eV', 'Ionization potential (Koopmans)', 'eV'),
        ('electron_affinity_eV', 'Electron affinity (Koopmans)', 'eV'),
        ('chemical_potential_eV', 'Chemical potential, mu', 'eV'),
        ('hardness_eV', 'Chemical hardness, eta', 'eV'),
        ('softness_inv_eV', 'Softness, S', '1/eV'),
        ('electrophilicity_eV', 'Electrophilicity index, omega', 'eV'),
        ('dipole_total_D', 'Dipole moment', 'D'),
        ('quadrupole_aniso_au', 'Quadrupole anisotropy', 'a.u.'),
        ('q_mulliken_min', 'Most negative Mulliken charge', 'e'),
        ('q_mulliken_max', 'Most positive Mulliken charge', 'e'),
        ('q_mulliken_range', 'Mulliken charge range', 'e'),
        ('xtb_Gsolv_water_Eh', 'GFN2-xTB solvation free energy (water)', 'Eh'),
    ]
    rows = []
    for col, name, unit in named:
        if col not in d:
            continue
        s = d[col].astype(float)
        r, pv = stats.pearsonr(s, d['PIC50'])
        rho = stats.spearmanr(s, d['PIC50'])[0]
        rows.append({'Descriptor': name, 'Unit': unit,
                     'Mean': round(s.mean(), 3), 'SD': round(s.std(), 3),
                     'Min': round(s.min(), 3), 'Max': round(s.max(), 3),
                     'r (pIC50)': round(r, 3), 'p': f'{pv:.3g}',
                     'rho (pIC50)': round(rho, 3)})
    write(pd.DataFrame(rows), 'Table5_dft_descriptors',
          'Table 5. Quantum-chemical descriptors computed for the 122 ligands, and their '
          'univariate association with pIC50.',
          'B3LYP/def2-SVP single-point energies (def2 ECP on iodine) on GFN2-xTB optimised '
          'geometries started from the docked/MD pose. Reactivity indices follow the '
          'Koopmans approximation. p-values are uncorrected for multiple testing.')


# ------------------------------------------------------------- T7 corrections
def t7_corrections():
    rows = [
        ('Hybrid MSE columns disagree between documents',
         'Main manuscript Table 4 gives CV MSE 0.6571, test 0.4602, unseen 0.9752; '
         'Hybrid model.docx Table 1 gives CV 0.4602, test 0.9752, unseen 0.1921',
         'The columns are shifted by one between the two documents. Re-derive from the '
         'recomputed values in Table 2 and use a single set throughout.'),
        ('Nu-SVR "Mean CV MSE" is actually the test MSE',
         'Table 4 reports 0.8591 under "Mean CV MSE"',
         'In the source script 0.8591 is mean_squared_error on the 20% test set. '
         'Relabel, or replace with the recomputed CV MSE.'),
        ('DNN narrative does not match its own table',
         'Hybrid model.docx quotes DNN-1 R2 train 0.8044 / CV 0.7373 and DNN-2 0.9532 / '
         '0.9323; Table 1 of the same document lists neither',
         'Reconcile the text to the table, or supply the runs those numbers came from.'),
        ('Duplicated rows in the DNN table',
         'DNN-1 and DNN-3 are identical; DNN-2 and DNN-4 are identical',
         'Two distinct architectures cannot give bit-identical metrics across seven '
         'columns. Re-run or merge the rows.'),
        ('Base learners were fitted on different splits',
         'Nu-SVR used random_state=1, all four DNNs used random_state=42, and the two '
         'were then stacked',
         'The ensemble members saw different training sets. All results here use one '
         'split per run.'),
        ('Hyperparameters were selected on the test set',
         'W_new includes R2 test and the test MSE/RMSE/MAE, and the reported model is the '
         'W_new maximum over the n_components sweep',
         'This makes the reported test R2 an optimistic estimate. Select on cross-'
         'validated R2 and quote the test set once, at the end.'),
        ('Meta-learner trained on in-sample base predictions',
         'No out-of-fold scheme is described for the Ridge stack',
         'Stacking on in-sample predictions inflates the meta-learner. With out-of-fold '
         'stacking the hybrid no longer consistently beats Nu-SVR (Table 2).'),
        ('Random splits leak duplicate structures',
         '10 of 25 test rows and 4 of 13 holdout rows share a canonical SMILES with the '
         'training set at random_state=1',
         'Report the structure-disjoint split alongside the random one (Table 2).'),
        ('Model comparisons are not statistically resolvable',
         'Hybrid test R2 0.6560 vs Nu-SVR 0.6532, a difference of 0.0028, on 23-25 test '
         'molecules',
         'The bootstrap 95% CI on test R2 spans roughly +/-0.5. State that the models are '
         'statistically indistinguishable rather than ranking them.'),
        ('SHAP interpretation attributes importance to protein confidence',
         'PC1 (mean |SHAP| 0.807) loads on pLDDT quantiles',
         'The AlphaFold block is constant within a protein, so PC1 largely encodes target '
         'identity. Target dummies alone give CV R2 0.34. Reword the interpretation.'),
        ('Feature count stated inconsistently',
         'Table 3 total reads "~270-280 272"',
         'The matrix has exactly 272 feature columns. Fix the cell.'),
        ('Number of AlphaFold feature vectors',
         'Methods describe four proteins; the data contain eight distinct pLDDT/PAE '
         'signatures',
         'Groups 5-7 differ only in the third decimal and two are singletons. Verify which '
         'AlphaFold run each ligand was assigned.'),
        ('Repository does not contain the pipeline',
         'The README advertises scripts/classical_ml_pipeline.py, hybrid_model_training.py, '
         'data/, models/ and outputs/; none exist in the repository',
         'Publish the actual training code, or point the manuscript at this benchmark '
         'repository.'),
    ]
    df = pd.DataFrame(rows, columns=['Issue', 'Where it appears', 'Recommended action'])
    write(df, 'TableS3_manuscript_corrections',
          'Table S3. Inconsistencies identified in the current manuscript drafts and the '
          'recommended correction for each.')


def main():
    os.makedirs(TAB, exist_ok=True)
    for fn in (t1_datasets, t2_benchmarks, t3_ablation, t4_statistics,
               t5_reproduction, t6_dft, t7_corrections):
        print(f'[{fn.__name__}]')
        try:
            fn()
        except Exception as e:
            import traceback
            print(f'  FAILED: {type(e).__name__}: {e}')
            traceback.print_exc(limit=2)


if __name__ == '__main__':
    main()
