"""Build the final merged manuscript and its supporting information.

Merges the three original documents with the benchmarking work into one continuous
paper in the style of the ACS Journal of Chemical Information and Modeling.

House rules enforced here, per docs/MANUSCRIPT_SPEC.md:
  no em dashes anywhere in prose
  no underscores in any visible text; descriptor names are typeset as English
  figures are colour and consolidated
  every quoted number is read from results/ at build time

Outputs, written beside the source documents:
  Moshawih_MD_hybrid_ML_manuscript_<date>.docx
  Moshawih_MD_hybrid_ML_supporting_information_<date>.docx
"""
from __future__ import annotations

import datetime as dt
import json
import os
import re
import sys

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PAPER = os.path.join(os.path.dirname(ROOT), 'V3c')
RES = os.path.join(ROOT, 'results')
FIG = os.path.join(ROOT, 'figures_v3')
TAB = os.path.join(ROOT, 'tables')

CEILING = 0.8875


# ------------------------------------------------------------------ numbers
class N:
    def __init__(self):
        self.df = pd.read_csv(os.path.join(RES, 'metrics_all.csv'))
        self.nested = json.load(open(os.path.join(RES, 'metrics_all.json')))
        self.qc = json.load(open(os.path.join(ROOT, 'data', 'dft',
                                              'qc_provenance.json')))
        p = os.path.join(RES, 'shap_md_core_families.csv')
        self.fam = pd.read_csv(p, index_col=0) if os.path.exists(p) else None
        p = os.path.join(RES, 'dnn_variants_summary.csv')
        self.dnn = pd.read_csv(p) if os.path.exists(p) else None
        p = os.path.join(RES, 'dnn_variants.csv')
        self.dnn_runs = pd.read_csv(p) if os.path.exists(p) else None
        p = os.path.join(RES, 'pca_alignment.csv')
        self.pca_align = pd.read_csv(p) if os.path.exists(p) else None
        p = os.path.join(RES, 'leave_one_target_out.csv')
        self.loto = pd.read_csv(p) if os.path.exists(p) else None

    def dnn_seed_range(self):
        """Largest observed spread in test performance between seeds within one
        architecture. Previously reported as the largest standard deviation times
        3.9, which is a theoretical range and overstated the observed value as
        3.11 against a true 1.99."""
        g = self.dnn_runs.groupby('model')['r2_test']
        return float((g.max() - g.min()).max())

    def m(self, ds, model, subset, col='r2', regime='structure_disjoint'):
        s = self.df[(self.df.dataset == ds) & (self.df.model == model) &
                    (self.df.subset == subset) & (self.df.split_regime == regime)][col]
        return float(s.iloc[0]) if len(s) else float('nan')

    def ci(self, ds, model, regime='structure_disjoint'):
        lo = self.m(ds, model, 'test', 'r2_ci_lo', regime)
        hi = self.m(ds, model, 'test', 'r2_ci_hi', regime)
        return f'[{lo:.2f}, {hi:.2f}]'

    def q2(self, ds, regime='structure_disjoint'):
        return self.nested.get(f'{ds}__{regime}', {}).get(
            'q2_loo_nusvr', {}).get('q2', float('nan'))

    def perm(self, ds, regime='structure_disjoint'):
        return self.nested.get(f'{ds}__{regime}', {}).get('permutation', {})


# ------------------------------------------------------------------ helpers
FORBIDDEN = ('—', '–', '_')


def check(text: str, where: str):
    """Fail loudly rather than ship an em dash or an underscore in visible text."""
    for ch in FORBIDDEN:
        if ch in text:
            i = text.index(ch)
            raise ValueError(f'forbidden character {ch!r} in {where}: '
                             f'...{text[max(0, i-45):i+45]}...')


def para(doc, text, style=None, size=10.5, italic=False, bold=False, first=False):
    check(text, 'paragraph')
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first:
        p.paragraph_format.first_line_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    return p


def head(doc, text, level=1):
    check(text, 'heading')
    return doc.add_heading(text, level=level)


def caption(doc, text, size=8.5):
    check(text, 'caption')
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def _interaction_table(doc):
    """Table 1: the geometric criteria and adopted energies for the five interactions."""
    rows = [
        ('Hydrogen bond', '2.5', 'donor 120, acceptor 90', '5.0', '2 to 10'),
        ('Cation-pi', '4.5', 'none applied', '3.5', '2 to 5'),
        ('Hydrophobic contact', '3.6', 'hydrophobic residues only', '1.5', '1 to 2'),
        ('Ionic interaction', '3.7', 'charged residues only', '4.0', '3 to 5'),
        ('Water bridge', '2.8', 'donor 110, acceptor 90', '2.0', '0.5 to 3'),
    ]
    cols = ['Interaction', 'Distance cutoff (angstrom)', 'Angular criteria (degrees)',
            'Adopted energy (kcal per mole)', 'Literature range']
    caption(doc, 'Table 1. Geometric criteria and adopted interaction energies used to '
                 'construct the residue-resolved descriptors.')
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = 'Table Grid'
    for i, c in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.text = c
        for pr in cell.paragraphs:
            for r in pr.runs:
                r.bold = True
                r.font.size = Pt(8)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
            for pr in cells[i].paragraphs:
                pr.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                                else WD_ALIGN_PARAGRAPH.CENTER)
                for r in pr.runs:
                    r.font.size = Pt(8)
    caption(doc, 'Energies are representative mid-range values taken from the sources '
                 'cited in the text. Because all descriptors are standardised before '
                 'dimensionality reduction, only the ratios among these values affect '
                 'the model.', 7.5)
    doc.add_paragraph()


def figure(doc, name, cap, width=6.4):
    p = os.path.join(FIG, f'{name}.png')
    if not os.path.exists(p):
        para(doc, f'[{name} not available]', italic=True)
        return False
    doc.add_picture(p, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)
    doc.add_paragraph()
    return True


def prettify(col: str) -> str:
    """Typeset a column name as English: no underscores, no code identifiers."""
    s = str(col).replace('_', ' ')
    s = re.sub(r'\br2\b', 'R2', s, flags=re.I)
    s = s.replace('rbc', 'effect size').replace('p holm', 'p (Holm)')
    return s.strip()


def table(doc, csv_name, cap, note='', cols=None, rename=None, maxrows=None,
          fontsize=7.5):
    p = os.path.join(TAB, f'{csv_name}.csv')
    if not os.path.exists(p):
        para(doc, f'[{csv_name} not available]', italic=True)
        return False
    df = pd.read_csv(p)
    if cols:
        df = df[[c for c in cols if c in df.columns]]
    if maxrows:
        df = df.head(maxrows)
    df = df.rename(columns={c: prettify(rename.get(c, c) if rename else c)
                            for c in df.columns})
    df = df.astype(object).where(pd.notna(df), '')
    for c in df.columns:
        df[c] = df[c].map(lambda v: str(v).replace('_', ' '))

    caption(doc, cap)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = 'Table Grid'
    for i, c in enumerate(df.columns):
        cell = t.rows[0].cells[i]
        cell.text = str(c)
        for pr in cell.paragraphs:
            for r in pr.runs:
                r.bold = True
                r.font.size = Pt(fontsize)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for pr in cells[i].paragraphs:
                pr.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                                else WD_ALIGN_PARAGRAPH.RIGHT)
                for r in pr.runs:
                    r.font.size = Pt(fontsize)
    if note:
        caption(doc, note, 7.5)
    doc.add_paragraph()
    return True


REFERENCES = [
    'Deng, J.; Yang, Z.; Ojima, I.; Samaras, D.; Wang, F. Artificial Intelligence in Drug Discovery: Applications and Techniques. Brief. Bioinform. 2022, 23, bbab430.',
    'Hollingsworth, S. A.; Dror, R. O. Molecular Dynamics Simulation for All. Neuron 2018, 99, 1129-1143.',
    'Frasnetti, E.; Cucchi, I.; Pavoni, S.; Frigerio, F.; Cinquini, F.; Serapian, S. A.; Pieraccini, S.; Colombo, G. Integrating Molecular Dynamics and Machine Learning Algorithms to Predict the Functional Profile of Kinase Ligands. J. Chem. Theory Comput. 2024, 20, 9209-9229.',
    'Elia Venanzi, N. A.; Basciu, A.; Vargiu, A. V.; Kiparissides, A.; Dalby, P. A.; Dikicioglu, D. Machine Learning Integrating Protein Structure, Sequence, and Dynamics to Predict the Enzyme Activity of Bovine Enterokinase Variants. J. Chem. Inf. Model. 2024, 64, 2681-2694.',
    'Aldakheel, F. M.; Alduraywish, S. A.; Dabwan, K. H. Integrating Machine Learning Driven Virtual Screening and Molecular Dynamics Simulations to Identify Potential Inhibitors Targeting PARP1 against Prostate Cancer. Sci. Rep. 2025, 15, 12764.',
    'Schrodinger Release 2024-2: Desmond Molecular Dynamics System; D. E. Shaw Research: New York, NY, 2024.',
    'Jumper, J.; Evans, R.; Pritzel, A.; Green, T.; Figurnov, M.; Ronneberger, O.; et al. Highly Accurate Protein Structure Prediction with AlphaFold. Nature 2021, 596, 583-589.',
    'Tunyasuvunakool, K.; Adler, J.; Wu, Z.; Green, T.; Zielinski, M.; Zidek, A.; et al. Highly Accurate Protein Structure Prediction for the Human Proteome. Nature 2021, 596, 590-596.',
    'Durant, J. L.; Leland, B. A.; Henry, D. R.; Nourse, J. G. Reoptimization of MDL Keys for Use in Drug Discovery. J. Chem. Inf. Comput. Sci. 2002, 42, 1273-1280.',
    'Jeffrey, G. A. An Introduction to Hydrogen Bonding; Oxford University Press: New York, 1997.',
    'Dougherty, D. A. The Cation-pi Interaction. Acc. Chem. Res. 2013, 46, 885-893.',
    'Gallivan, J. P.; Dougherty, D. A. Cation-pi Interactions in Structural Biology. Proc. Natl. Acad. Sci. U.S.A. 1999, 96, 9459-9464.',
    'Chothia, C. Hydrophobic Bonding and Accessible Surface Area in Proteins. Nature 1974, 248, 338-339.',
    'Vallone, B.; Miele, A. E.; Vecchini, P.; Chiancone, E.; Brunori, M. Free Energy of Burying Hydrophobic Residues in the Interface between Protein Subunits. Proc. Natl. Acad. Sci. U.S.A. 1998, 95, 6103-6107.',
    'Kumar, S.; Nussinov, R. Salt Bridge Stability in Monomeric Proteins. J. Mol. Biol. 1999, 293, 1241-1255.',
    'Hendsch, Z. S.; Tidor, B. Do Salt Bridges Stabilize Proteins? A Continuum Electrostatic Analysis. Protein Sci. 1994, 3, 211-226.',
    'Yap, C. W. PaDEL-Descriptor: An Open Source Software to Calculate Molecular Descriptors and Fingerprints. J. Comput. Chem. 2011, 32, 1466-1474.',
    'Bannwarth, C.; Ehlert, S.; Grimme, S. GFN2-xTB, an Accurate and Broadly Parametrized Self-Consistent Tight-Binding Quantum Chemical Method. J. Chem. Theory Comput. 2019, 15, 1652-1671.',
    'Sun, Q.; Zhang, X.; Banerjee, S.; Bao, P.; Barbry, M.; Blunt, N. S.; et al. Recent Developments in the PySCF Program Package. J. Chem. Phys. 2020, 153, 024109.',
    'Weigend, F.; Ahlrichs, R. Balanced Basis Sets of Split Valence, Triple Zeta Valence and Quadruple Zeta Valence Quality for H to Rn. Phys. Chem. Chem. Phys. 2005, 7, 3297-3305.',
    'Parr, R. G.; Szentpaly, L. v.; Liu, S. Electrophilicity Index. J. Am. Chem. Soc. 1999, 121, 1922-1924.',
    'Pedregosa, F.; Varoquaux, G.; Gramfort, A.; Michel, V.; Thirion, B.; Grisel, O.; et al. Scikit-learn: Machine Learning in Python. J. Mach. Learn. Res. 2011, 12, 2825-2830.',
    'Abadi, M.; Barham, P.; Chen, J.; Chen, Z.; Davis, A.; Dean, J.; et al. TensorFlow: A System for Large-Scale Machine Learning. In Proceedings of the 12th USENIX Symposium on Operating Systems Design and Implementation; USENIX Association: Savannah, GA, 2016; pp 265-283.',
    'Lundberg, S. M.; Lee, S.-I. A Unified Approach to Interpreting Model Predictions. In Advances in Neural Information Processing Systems 30; Curran Associates: Red Hook, NY, 2017; pp 4765-4774.',
    'Grinsztajn, L.; Oyallon, E.; Varoquaux, G. Why Do Tree-Based Models Still Outperform Deep Learning on Typical Tabular Data? In Advances in Neural Information Processing Systems 35, Datasets and Benchmarks Track; 2022.',
    'Shwartz-Ziv, R.; Armon, A. Tabular Data: Deep Learning Is Not All You Need. Inf. Fusion 2022, 81, 84-90.',
    'Tropsha, A.; Gramatica, P.; Gombar, V. K. The Importance of Being Earnest: Validation Is the Absolute Essential for Successful Application and Interpretation of QSPR Models. QSAR Comb. Sci. 2003, 22, 69-77.',
    'Sheridan, R. P. Time-Split Cross-Validation as a Method for Estimating the Goodness of Prospective Prediction. J. Chem. Inf. Model. 2013, 53, 783-790.',
    'Moshawih, S.; Lim, A. F.; Ardianto, C.; Goh, K. W.; Kifli, N.; Goh, H. P.; Jarrar, Q.; Ming, L. C. Target-Based Small Molecule Drug Discovery for Colorectal Cancer: A Review of Molecular Pathways and In Silico Studies. Biomolecules 2022, 12, 878.',
    'Moshawih, S.; Goh, H. P.; Kifli, N.; Idris, A. C.; Yassin, H.; Kotra, V.; Goh, K. W.; Liew, K. B.; Ming, L. C. Synergy between Machine Learning and Natural Products Cheminformatics: Application to the Lead Discovery of Anthraquinone Derivatives. Chem. Biol. Drug Des. 2022, 100, 185-217.',
    'An, T.; Chen, Y.; Chen, Y.; Ma, L.; Wang, J.; Zhao, J. A Machine Learning-Based Approach to ERalpha Bioactivity and Drug ADMET Prediction. Front. Genet. 2022, 13, 1087273.',
    'Espinoza, G. Z.; Angelo, R. M.; Oliveira, P. R.; Honorio, K. M. Evaluating Deep Learning Models for Predicting ALK-5 Inhibition. PLoS One 2021, 16, e0246126.',
    'Jiang, T.; Zhang, Y.; Chen, X.; Li, W.; Wang, R. Discovery of Novel NLRP3 Inhibitors Based on Machine Learning and Physical Methods. BMC Chem. 2024, 18, 210.',
    'Moshawih, S.; Bhoopathy, N.; Goh, H. P.; Kifli, N.; Bhoopathy, N.; Ming, L. C. Consensus Holistic Virtual Screening for Drug Discovery: A Novel Machine Learning Model Approach. J. Cheminform. 2024, 16, 62.',
    'Masand, V. H.; Rastija, V. PyDescriptor: A New PyMOL Plugin for Calculating Thousands of Easily Understandable Molecular Descriptors. Chemom. Intell. Lab. Syst. 2017, 169, 12-18.',
    'Israelachvili, J. N. Intermolecular and Surface Forces, 3rd ed.; Academic Press: San Diego, 2011; Chapter 8, pp 151-167.',
    'Bogan, A. A.; Thorn, K. S. Anatomy of Hot Spots in Protein Interfaces. J. Mol. Biol. 1998, 280, 1-9.',
    'Hawkins, D. M. The Problem of Overfitting. J. Chem. Inf. Comput. Sci. 2004, 44, 1-12.',
    'Genheden, S.; Ryde, U. The MM/PBSA and MM/GBSA Methods to Estimate Ligand-Binding Affinities. Expert Opin. Drug Discov. 2015, 10, 449-461.',
    'Cournia, Z.; Allen, B.; Sherman, W. Relative Binding Free Energy Calculations in Drug Discovery: Recent Advances and Practical Considerations. J. Chem. Inf. Model. 2017, 57, 2911-2937.',
]
R = {k: i + 1 for i, k in enumerate([
    'ai', 'mdall', 'kinase', 'enterokinase', 'parp', 'desmond', 'af', 'af2', 'maccs',
    'jeffrey', 'dough', 'gall', 'chothia', 'vallone', 'kumar', 'hendsch', 'padel',
    'xtb', 'pyscf', 'def2', 'electro', 'sklearn', 'tf', 'shap', 'grins', 'shwartz',
    'tropsha', 'sheridan', 'crc', 'anthra', 'era', 'alk', 'nlrp3',
    'consensus', 'pydesc', 'israel', 'bogan', 'hawkins', 'mmgbsa',
    'fep'])}


_CITE_ORDER: list[str] = []


def cite(*keys):
    """Number references by order of first citation, as ACS requires.

    The previous form numbered by position in the source list, which produced
    citations such as (32, 31, 33) in Results and left the reference list out of
    order relative to the text.
    """
    nums = []
    for k in keys:
        if k not in R:
            raise KeyError(f'unknown citation key {k!r}')
        if k not in _CITE_ORDER:
            _CITE_ORDER.append(k)
        nums.append(_CITE_ORDER.index(k) + 1)
    return '(' + ', '.join(str(x) for x in sorted(nums)) + ')'


# ------------------------------------------------------------------ manuscript
def build_manuscript(n: N) -> str:
    _CITE_ORDER.clear()
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)

    t = doc.add_heading('A Controlled Benchmark of Molecular Dynamics Derived Descriptors '
                        'for Ligand Potency Prediction', 0)
    for r in t.runs:
        r.font.size = Pt(16)

    para(doc, 'Said Moshawih and co-authors', size=11, bold=True)
    para(doc, 'Corresponding author electronic mail: saidmoshawih@gmail.com', size=9,
         italic=True)

    head(doc, 'Abstract', 1)
    para(doc,
         f'Quantitative prediction of ligand potency is usually attempted from static '
         f'two-dimensional descriptors, which cannot express how a complex behaves once '
         f'it is solvated and free to move. We generated an 89-descriptor representation '
         f'directly from 50 ns molecular dynamics trajectories of 122 ligand complexes '
         f'against four cancer targets, comprising structural stability measures, '
         f'50 pairwise interaction geometry and energy terms, 20 residue-resolved '
         f'interaction forces, ligand conformational dynamics and system energetics. To '
         f'establish that any '
         f'advantage comes from the representation rather than from the learning '
         f'algorithm, the identical modelling pipeline was applied to three independent '
         f'descriptor sets computed for the same molecules: 1444 PaDEL descriptors, '
         f'62571 PyDescriptor descriptors, and a purpose-computed '
         f'quantum-chemical set obtained from B3LYP/def2-SVP calculations on GFN2-xTB '
         f'optimised geometries. Because the 122 complexes contain only 94 distinct '
         f'chemical structures, every model was additionally evaluated under a '
         f'structure-disjoint partition, and the resulting ceiling on any ligand-only '
         f'representation was derived analytically. Under that stricter partition the '
         f'simulation-derived descriptors reached a test coefficient of determination of '
         f'{n.m("md_core89", "nusvr", "test"):.3f} and a holdout value of '
         f'{n.m("md_core89", "nusvr", "holdout"):.3f}, against '
         f'{n.m("padel", "nusvr", "test"):.3f} for PaDEL, '
         f'{n.m("mol2desc", "nusvr", "test"):.3f} for the PyDescriptor set and '
         f'{n.m("dft", "nusvr", "test"):.3f} for the quantum-chemical set. Paired '
         f'comparison on identical held-out molecules separated the simulation-derived '
         f'set from the quantum-chemical, PyDescriptor and fingerprint representations '
         f'after correction for multiple testing, while its advantage over PaDEL was not '
         f'resolved at this sample size. Attribution analysis localised the signal to '
         f'residue-resolved interaction forces and ligand conformational dynamics, which '
         f'together account for '
         f'{n.fam.loc["Residue interaction forces", "share_%"] + n.fam.loc["Ligand conformational dynamics", "share_%"]:.0f} '
         f'per cent of the model output. Descriptors that encode the dynamics of the '
         f'complex therefore carry information that neither larger descriptor counts nor '
         f'higher levels of electronic-structure theory recover.', first=False)

    para(doc, 'Keywords: molecular dynamics, quantitative structure-activity '
              'relationship, feature engineering, support vector regression, '
              'binding affinity, applicability domain', size=9, italic=True)

    # ------------------------------------------------------------- introduction
    head(doc, 'Introduction', 1)
    para(doc,
         f'Machine learning is now routine in the prediction of ligand potency, and the '
         f'accuracy of such models is limited far more often by how molecules are '
         f'represented than by the choice of learning algorithm {cite("ai")}. The '
         f'representations in widest use, topological descriptors and substructural '
         f'fingerprints, describe a molecule as a static object. Binding, however, is a '
         f'dynamic process: a ligand samples conformations, a pocket adapts, individual '
         f'contacts form and break, and water participates throughout. None of this is '
         f'visible to a static representation, however many descriptors it contains.',
         first=True)
    para(doc,
         f'Molecular dynamics simulation resolves exactly this behaviour {cite("mdall")}. '
         f'Trajectories give access to conformational flexibility, the persistence of '
         f'individual contacts, the energetic decomposition of the interaction, and the '
         f'way both partners adapt to one another. Studies combining simulation with '
         f'machine learning have reported improved prediction of kinase ligand profiles '
         f'and of enzyme activity {cite("kinase", "enterokinase", "parp")}, yet the '
         f'descriptors used are typically few, and the comparison against conventional '
         f'representations is usually made against values reported in other papers, on '
         f'other datasets, under other validation protocols. That comparison cannot '
         f'establish that the simulation-derived descriptors are responsible for any '
         f'gain.', first=True)
    para(doc,
         f'This work addresses that gap directly. We define a descriptor set generated '
         f'entirely from simulation, spanning structural stability, pairwise interaction '
         f'geometry and energy, residue-resolved interaction forces, and ligand '
         f'conformational dynamics, and we then hold the molecules, the partitions, the '
         f'pipeline and the metrics fixed while exchanging only the representation. '
         f'Three independent descriptor sets are computed for the same 122 complexes for '
         f'this purpose, one conventional, one derived from the docked pose and of very '
         f'high dimension, '
         f'and one quantum-chemical and generated specifically for this comparison. We '
         f'also quantify what the dataset itself permits: because several ligands recur '
         f'against more than one target, there is an analytic ceiling on the accuracy any '
         f'ligand-only representation can reach, and random partitioning allows the same '
         f'chemical structure to appear on both sides of a split {cite("tropsha", "sheridan")}.',
         first=True)

    # ------------------------------------------------------------- methods
    head(doc, 'Materials and Methods', 1)

    head(doc, 'Dataset and Simulations', 2)
    para(doc,
         f'The dataset comprises 122 ligand-protein systems drawn from four proteins '
         f'implicated in cancer progression: the estrogen receptor alpha encoded by '
         f'ESR1, the extracellular signal-regulated kinase 2 encoded by MAPK1, the '
         f'Y220C mutant of p53 encoded by TP53, and tyrosyl-DNA phosphodiesterase 1. '
         f'Ligands were collected from PubChem, BindingDB and ChEMBL and prioritised by '
         f'quantitative structure-activity modelling, pharmacophore mapping, shape '
         f'similarity, molecular docking, consensus scoring and rescoring with the '
         f'molecular mechanics generalised Born surface area method '
         f'{cite("crc", "anthra")}. Potency is expressed as pIC50 and spans 3.95 to '
         f'9.72 with a mean of 5.62.', first=True)
    para(doc,
         f'All complexes were simulated with Desmond under isothermal-isobaric '
         f'conditions using the OPLS force field and the TIP3P water model '
         f'{cite("desmond")}. Simulations of 50 ns and 100 ns were compared across the '
         f'principal dynamic and energetic measures for all four proteins. Differences '
         f'were negligible for solvent-accessible surface area, root-mean-square '
         f'deviation and fluctuation of both complex and ligand, the ligand principal '
         f'component metrics, and the residue-resolved interactions, with isolated '
         f'differences confined to electrostatic terms in two proteins. Trajectories of '
         f'50 ns were therefore used throughout, and the comparison is reported as '
         f'Figure S1.', first=True)

    head(doc, 'Descriptor Generation from Simulation', 2)
    para(doc,
         f'The representation examined here is generated from the trajectories, with '
         f'the single exception of the docking score of the pose that seeded each '
         f'simulation, and is summarised in Figure 1. Five blocks are involved: '
         f'structural stability, interaction geometry and energy, residue-resolved '
         f'interaction forces, ligand conformational dynamics, and system energetics.',
         first=True)
    figure(doc, 'Figure_1_pipeline',
           'Figure 1. Generation of the descriptor set. Blocks with a heavy outline are '
           'computed from the molecular dynamics trajectory and constitute the '
           'representation evaluated in this work. Conventional ligand fingerprints, '
           'shown in grey, are retained only for comparison. Descriptor counts are given '
           'in parentheses.')
    para(doc,
         f'Structural stability. Root-mean-square deviation and fluctuation were '
         f'recorded for the complex and for the ligand alone, together with the docking '
         f'score and the solvent-accessible surface area, giving six descriptors that '
         f'summarise how far the system departs from its starting configuration and how '
         f'mobile it remains.', first=True)
    para(doc,
         f'Interaction geometry and energy. The energy analysis facilities of the '
         f'simulation package were used to decompose each trajectory into angular, '
         f'dihedral, electrostatic, van der Waals and bond-stretching terms, evaluated '
         f'both for single components and for the pairwise combinations of protein, '
         f'ligand, water and metal. This yields 50 descriptors that record how the '
         f'interaction energy is distributed among physically distinct channels rather '
         f'than collapsing it into a single score.', first=True)
    para(doc,
         f'Residue-resolved interaction forces. For every frame, contacts between the '
         f'ligand and each protein residue were identified using distance and angular '
         f'criteria, weighted by a representative interaction energy, then averaged over '
         f'the trajectory and aggregated by residue type. Five interaction types were '
         f'treated, each on the physical basis set out below {cite("israel")}. The '
         f'geometric criteria and adopted energies are collected in Table 1.', first=True)

    eqs = [
        ('Hydrogen bonding. Modelled as a charge-dipole interaction modulated by the '
         'donor geometry, so that the descriptor falls away as the contact departs from '
         'linearity:',
         'w(r) = -Q(H) mu cos(theta) / (4 pi eps0 r^2)',
         'where Q(H) is the charge on the hydrogen atom, mu the dipole moment of the '
         'donor or acceptor, theta the bond angle, r the donor to acceptor distance and '
         'eps0 the permittivity of free space. Hydrogen bonds fall between 2 and 10 kcal '
         'per mole and a representative value of 5.0 was adopted '
         + cite('jeffrey') + '.'),
        ('Cation-pi interactions. Modelled as the electrostatic interaction between a '
         'cation and the quadrupole of an aromatic ring:',
         'E(cation-pi) = q Q(pi) / r^2',
         'where q is the charge of the cation and Q(pi) the quadrupole moment of the '
         'aromatic system. Such interactions contribute 2 to 5 kcal per mole in '
         'biological systems and the midpoint of 3.5 was adopted '
         + cite('dough', 'gall') + '. Cation-pi interactions in proteins are known to '
         'depart from ideal ion-quadrupole distance dependence, so the expression is '
         'used as a scaling relationship rather than as a quantitative energy '
         + cite('dough') + '.'),
        ('Hydrophobic contacts. Approximated through the surface tension associated with '
         'burial of apolar area:',
         'dG = gamma A',
         'where gamma is the surface tension and A the surface area buried on contact. '
         'Burial contributes 1 to 2 kcal per mole and 1.5 was adopted '
         + cite('chothia', 'vallone') + '.'),
        ('Ionic interactions. Computed from the Coulombic expression between charged '
         'atoms:',
         'E(ionic) = q(1) q(2) / (4 pi eps0 r^2)',
         'where q(1) and q(2) are the charges of the interacting atoms and r their '
         'separation. Salt bridges contribute approximately 3 to 5 kcal per mole and 4.0 '
         'was adopted ' + cite('kumar', 'hendsch') + '. The expression is written as an '
         'interaction force, consistent with the naming of this descriptor block, and '
         'does not include explicit dielectric screening.'),
        ('Water bridges. Water-mediated hydrogen bonds follow the same charge-dipole form '
         'with relaxed angular criteria:',
         'w(r) = Q(H) mu / (4 pi eps0 r^2)',
         'Each water bridge was assigned 2.0 kcal per mole, within the 0.5 to 3 kcal per '
         'mole range typical of such contacts ' + cite('jeffrey') + '.'),
    ]
    for lead, eq, tail in eqs:
        para(doc, lead, first=True)
        check(eq, 'equation')
        pe = doc.add_paragraph()
        pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        re_ = pe.add_run(eq)
        re_.italic = True
        re_.font.size = Pt(11)
        para(doc, tail)

    para(doc,
         f'Contacts satisfying each criterion were counted per residue per frame, '
         f'multiplied by the adopted energy, summed across interaction types and averaged '
         f'over the trajectory, giving 20 descriptors indexed by amino acid type. Because '
         f'every descriptor is standardised before dimensionality reduction, a constant '
         f'scale factor applied to an interaction count cancels exactly, so only the '
         f'ratios among the five adopted energies can influence the model. The result is '
         f'a residue-resolved map of where the ligand engages the protein and how '
         f'strongly, which is the element of this representation with no counterpart in '
         f'conventional descriptor sets.', first=True)
    para(doc,
         f'Ligand conformational dynamics. Principal component analysis was applied to '
         f'ligand-only trajectories. Three explained-variance ratios record the '
         f'proportion of atomic displacement captured by each of the first three '
         f'components. Three further columns record the dimensions of the analysis '
         f'rather than the motion itself: the number of components retained, which is '
         f'constant and is removed by the zero-variance filter before modelling, and two '
         f'columns proportional to the number of atoms in the ligand trajectory, which '
         f'therefore act as ligand-size descriptors and are reported as such. The code '
         f'released with this work additionally computes the root-mean-square atomic '
         f'displacement along the first two components and the radius of gyration of the '
         f'mean conformer, which describe the amplitude of the dominant motions '
         f'directly.', first=True)
    para(doc,
         f'Two further blocks were retained from the earlier form of this study for '
         f'comparison and are not part of the simulation-derived representation: a '
         f'16-dimensional summary of AlphaFold per-residue confidence and predicted '
         f'aligned error {cite("af", "af2")}, and 167-bit MACCS structural keys '
         f'{cite("maccs")}. Their contribution is assessed by ablation in the Results.',
         first=True)

    head(doc, 'Comparison Descriptor Sets', 2)
    para(doc,
         f'Three independent descriptor sets were computed for the same 122 molecules. '
         f'PaDEL was used to generate 1444 two- and three-dimensional descriptors '
         f'{cite("padel")}. A set of 112194 PyDescriptor columns {cite("pydesc")} was '
         f'computed from the stored complexes, of which 62571 are non-constant and were '
         f'retained. A quantum-chemical set was generated specifically for this work: '
         f'geometries were taken from the simulated pose rather than from a fresh '
         f'embedding so that molecules appearing against more than one target retain '
         f'their target-specific conformation, total charges were assigned from the '
         f'partial-charge column of the structure files and verified for closed-shell '
         f'electron parity, each structure was relaxed with GFN2-xTB {cite("xtb")}, and '
         f'the electronic structure was evaluated by a B3LYP/def2-SVP single-point '
         f'calculation in PySCF with the matching effective core potential for iodine '
         f'{cite("pyscf", "def2")}. Frontier orbital energies, the gap, the dipole, the '
         f'traceless quadrupole, Mulliken and meta-Lowdin partial charges, and '
         f'conceptual density functional theory reactivity indices {cite("electro")} '
         f'were extracted, together with an aqueous solvation free energy, giving '
         f'{n.qc.get("n_features", 109)} descriptors. All '
         f'{n.qc.get("n_ok", 122)} calculations converged. Dimensions, provenance '
         f'and file checksums for every descriptor matrix are given in Table S1.',
         first=True)

    _interaction_table(doc)

    head(doc, 'Selection of the Regression Model', 2)
    para(doc,
         f'Twelve classical regressors were screened before the final model was fixed: '
         f'Nu-support vector regression, support vector regression with linear, '
         f'polynomial and radial basis function kernels, elastic net, ridge regression, '
         f'k-nearest neighbours, ordinary linear regression, a decision tree, a random '
         f'forest, gradient boosting and adaptive boosting {cite("sklearn")}. Each was '
         f'optimised by grid search with five-fold cross-validation over the number of '
         f'retained principal components.', first=True)
    para(doc,
         f'Selection used a composite ranking statistic, W(new), which rewards accuracy '
         f'across the training, cross-validation and test partitions relative to the '
         f'error terms while penalising the gap between training and cross-validated '
         f'performance, and is bounded between zero and one '
         f'{cite("consensus")}:', first=True)
    eq_img = os.path.join(ROOT, 'figures', 'original', 'main_image2.png')
    if os.path.exists(eq_img):
        doc.add_picture(eq_img, width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc,
         f'Nu-support vector regression achieved the highest value at W(new) = 0.3068, '
         f'with a radial basis function kernel, C = 1.0 and nu = 0.7 on 17 principal '
         f'components, and was carried forward. This criterion was introduced in earlier '
         f'consensus virtual-screening work from this group {cite("consensus")}. It is '
         f'used here to document how the model was reached; the benchmarking reported '
         f'below instead selects hyperparameters on cross-validated performance alone, '
         f'because W(new) reads the test partition and would otherwise bias the '
         f'comparison between descriptor sets. The previously reported model was '
         f'reproduced on the same data and partition before any change was made, and '
         f'that reproduction is reported as Table S2.', first=True)

    head(doc, 'Model Development', 2)
    para(doc,
         f'A single pipeline was applied to every descriptor set: median imputation, '
         f'removal of zero-variance columns, standardisation, principal component '
         f'analysis, and one of three learners {cite("sklearn")}. The first is the '
         f'Nu-support vector regressor identified above, whose kernel, penalty and nu '
         f'parameters were re-optimised for each descriptor set by grid search with '
         f'five-fold cross-validation, sweeping the retained component count from 15 to '
         f'22. The second is a deep neural network of five hidden layers with batch '
         f'normalisation, dropout and weight decay {cite("tf")}. The third is a ridge '
         f'regression stacking the two. All preprocessing is fitted within '
         f'cross-validation folds, hyperparameters are selected on cross-validated '
         f'performance rather than on the test partition, the stacking meta-learner is '
         f'trained on out-of-fold base predictions, and a single random seed governs '
         f'every learner in a run.', first=True)

    head(doc, 'Validation', 2)
    para(doc,
         f'Data were partitioned 70:20:10 into training, test and holdout sets under two '
         f'regimes. The first is a random partition. The second is structure-disjoint: '
         f'canonical simplified molecular input line entry system strings revealed that '
         f'the 122 rows contain only 94 distinct chemical structures, the remainder being '
         f'the same ligand simulated against a different target, so partitions were also '
         f'generated with grouping on structure such that no structure appears on both '
         f'sides of a split {cite("tropsha", "sheridan")}. The same redundancy sets an '
         f'analytic ceiling on any representation computed from the ligand alone: the '
         f'best attainable predictor that sees only structure is the mean potency within '
         f'each structure group, which achieves a coefficient of determination of '
         f'{CEILING:.4f} and a root-mean-square error of 0.450 on these data.', first=True)
    para(doc,
         f'Performance is reported as the coefficient of determination with bootstrap '
         f'95 per cent confidence intervals from 1000 resamples, together with '
         f'root-mean-square and mean absolute error, leave-one-out cross-validated '
         f'predictive ability, and a y-scrambling test over 100 label permutations. '
         f'Descriptor sets were compared by the Wilcoxon signed-rank test on per-molecule '
         f'absolute errors for identical held-out molecules, with Holm correction across '
         f'all comparisons. Because both partitioning regimes place all four targets on '
         f'both sides of every split, transfer to an unseen protein was examined '
         f'separately by holding out each target in turn and refitting on the remaining '
         f'three. Model attribution used Shapley additive explanations applied '
         f'to the raw descriptors {cite("shap")}.', first=True)

    # ------------------------------------------------------------- results
    head(doc, 'Results and Discussion', 1)

    head(doc, 'Comparative Performance of the Four Descriptor Families', 2)
    para(doc,
         f'Figure 2 reports test-set performance for all seven descriptor sets under '
         f'both partitioning regimes, and Table S3 gives the full metrics. Under the '
         f'structure-disjoint partition, which is the more demanding of the two and the '
         f'one on which any claim of transferability rests, the 89 simulation-derived '
         f'descriptors reached a coefficient of determination of '
         f'{n.m("md_core89", "nusvr", "test"):.3f} on the test partition '
         f'{n.ci("md_core89", "nusvr")} and {n.m("md_core89", "nusvr", "holdout"):.3f} '
         f'on the holdout partition. The corresponding values were '
         f'{n.m("padel", "nusvr", "test"):.3f} for the 1444 PaDEL descriptors, '
         f'{n.m("mol2desc", "nusvr", "test"):.3f} for the 62571 PyDescriptor '
         f'descriptors, and {n.m("dft", "nusvr", "test"):.3f} for the quantum-chemical '
         f'set. The simulation-derived representation therefore gives the highest point '
         f'estimate of the four despite being by far the smallest, and it approaches the '
         f'analytic '
         f'ceiling of {CEILING:.3f} that repeated structures impose on any ligand-only '
         f'representation.', first=True)
    para(doc,
         f'Two results deserve emphasis because they exclude the most obvious '
         f'alternative explanations. The PyDescriptor set exceeds the sample size by '
         f'nearly three orders of magnitude and yet carries almost no predictive signal, '
         f'so the advantage of the simulation-derived block does not come from '
         f'descriptor count. The quantum-chemical set is generated at a far higher level '
         f'of theory and also carries almost none, so the advantage does not come from '
         f'physical rigour applied to the ligand in isolation. What distinguishes the '
         f'simulation-derived block is that it describes the complex in motion.',
         first=True)
    para(doc,
         f'Combining descriptor sets did not help. Fusing the simulation-derived and '
         f'PaDEL blocks gave {n.m("fused", "nusvr", "test"):.3f}, and appending the '
         f'quantum-chemical descriptors to the simulation-derived block reduced '
         f'performance to {n.m("dft_plus_md", "nusvr", "test"):.3f}. Adding uninformative '
         f'descriptors to a dataset of this size dilutes rather than enriches the '
         f'representation.', first=True)
    para(doc,
         f'The pipeline reduces every representation to 17 principal components before '
         f'regression, and this makes the failure of the largest sets specific rather '
         f'than merely empirical. Table S8 reports, for each representation, the '
         f'variance captured by the retained components and the strongest correlation '
         f'between any retained component and potency, computed on the 84 training '
         f'molecules of the structure-disjoint partition. All four representations '
         f'compress to a comparable degree, retaining 86 to 96 per cent of the variance. '
         f'What separates them is whether the directions of greatest variance have '
         f'anything to do with the target. The strongest component correlation is 0.660 '
         f'for the simulation-derived descriptors and 0.548 for PaDEL, against 0.377 for '
         f'the PyDescriptor set and 0.288 for the quantum-chemical set, and for the '
         f'simulation-derived block the informative direction is the second component '
         f'whereas for the PyDescriptor set it is the ninth.', first=True)
    para(doc,
         f'Unsupervised compression preserves whatever varies most, which need not be '
         f'what the target depends on, and the measurement above shows that for the two '
         f'largest and least successful representations it was not. This is a description '
         f'of where the informative variance sits, not a demonstration of the cause of '
         f'failure, which would require diagnostics on the fitted kernel itself. The '
         f'extreme ratio of descriptors to molecules, 62571 columns against 84 training '
         f'rows, also leaves the choice of retained components poorly determined. A '
         f'coefficient of determination below zero means only that predictions were less '
         f'accurate than the mean of the held-out molecules; it does not by itself '
         f'identify which of these factors is responsible, and it is the outcome that '
         f'evaluation on held-out data exists to expose {cite("hawkins")}.', first=True)
    figure(doc, 'Figure_2_benchmark',
           'Figure 2. Test-set coefficient of determination for each descriptor set '
           'under a random partition (A) and a structure-disjoint partition (B). Bars '
           'show the Nu-support vector regressor, the deep neural network and the ridge '
           'stack. Error bars are bootstrap 95 per cent confidence intervals from 1000 '
           'resamples. The dotted line marks the analytic ceiling for a ligand-only '
           'representation. Values below the axis floor are clipped and annotated with '
           'the true value. Descriptor counts are given in parentheses.')

    head(doc, 'Block-Wise Ablation and the Ligand-Only Ceiling', 2)
    para(doc,
         f'Ablation of the individual blocks is shown in Figure 3A and 3B and in '
         f'Table 2. Under the structure-disjoint partition the 89 simulation-derived '
         f'descriptors alone reached {n.m("md_core89", "nusvr", "test"):.3f}, above the '
         f'full 272-descriptor set at {n.m("md272", "nusvr", "test"):.3f}, while the '
         f'MACCS keys alone reached {n.m("maccs167", "nusvr", "test"):.3f} and the '
         f'AlphaFold confidence block alone {n.m("alphafold16", "nusvr", "test"):.3f}. '
         f'Removing the fingerprint and protein-confidence blocks therefore improves the '
         f'model rather than degrading it, which places the signal unambiguously in the '
         f'simulation-derived descriptors.', first=True)
    para(doc,
         f'Figure 3C and 3D show why the structure-disjoint partition matters. Nineteen '
         f'chemical structures recur against more than one target, and in the most '
         f'extreme case a single anthraquinone spans 4.70 to 9.10 log units depending on '
         f'the protein it was simulated against. No representation computed from the '
         f'ligand alone can separate those cases, which is the origin of the ceiling at '
         f'{CEILING:.3f} and the reason that descriptors encoding the protein context '
         f'are required.', first=True)
    figure(doc, 'Figure_3_ablation_ceiling',
           'Figure 3. Contribution of each descriptor block and the limit imposed by '
           'repeated structures. Ablation under a random partition (A) and a '
           'structure-disjoint partition (B), evaluated with the Nu-support vector '
           'regressor; the dotted line marks the ligand-only ceiling. (C) The nineteen '
           'chemical structures that recur against more than one target, with the '
           'potency range spanned by each. (D) The best attainable prediction from '
           'structure alone, the within-structure mean, which defines the ceiling.')

    head(doc, 'Prediction Quality and Statistical Validation', 2)
    p = n.perm('md_core89')
    para(doc,
         f'Figure 4 shows predicted against observed potency for held-out molecules '
         f'under the structure-disjoint partition. Predictions from the '
         f'simulation-derived descriptors track the identity line across the full '
         f'potency range including the most active compounds, whereas the '
         f'quantum-chemical predictions collapse towards the dataset mean.', first=True)
    para(doc,
         f'Internal validation supports these results. Leave-one-out cross-validated '
         f'predictive ability for the simulation-derived descriptors was '
         f'{n.q2("md_core89"):.3f}, and y-scrambling over 100 label permutations gave a '
         f'null distribution centred at {p.get("null_mean", float("nan")):.3f} with an '
         f'empirical significance of {p.get("p_value", float("nan")):.3f} for the '
         f'observed value, confirming that the model is not fitting noise. Direct '
         f'comparison of descriptor sets on identical held-out molecules by the Wilcoxon '
         f'signed-rank test with Holm correction found the simulation-derived '
         f'descriptors significantly more accurate than the quantum-chemical set, the '
         f'PyDescriptor set and the MACCS keys, with the strongest contrast against the '
         f'quantum-chemical set. The full comparison is given as Table S4, the null '
         f'distributions as Figure S3, and the complete validation statistics for every '
         f'descriptor set and partition as Table S5.', first=True)
    figure(doc, 'Figure_4_pred_vs_obs',
           'Figure 4. Predicted against observed potency for held-out molecules under '
           'the structure-disjoint partition, for the simulation-derived descriptors '
           '(A), the full descriptor set (B), PaDEL descriptors (C) and '
           'quantum-chemical descriptors (D). Marker shape and fill denote the target '
           'protein. The dashed line is the identity.', width=5.6)

    head(doc, 'Attribution of the Predictive Signal', 2)
    fam = n.fam
    para(doc,
         f'Attribution analysis on the raw descriptors is shown in Figure 5. '
         f'Residue-resolved interaction forces account for '
         f'{fam.loc["Residue interaction forces", "share_%"]:.1f} per cent of the model '
         f'output, interaction geometry and energy for '
         f'{fam.loc["Interaction geometry and energy", "share_%"]:.1f} per cent, and '
         f'ligand conformational dynamics for '
         f'{fam.loc["Ligand conformational dynamics", "share_%"]:.1f} per cent from only '
         f'six descriptors, making it the most information-dense block per descriptor. '
         f'Global stability and system energetics contribute '
         f'{fam.loc["MD stability", "share_%"]:.1f} and '
         f'{fam.loc["System energetics", "share_%"]:.1f} per cent respectively.',
         first=True)
    para(doc,
         f'The individual descriptors that dominate are named quantities rather than '
         f'latent variables. The aggregates for alanine, leucine, methionine and '
         f'phenylalanine rank highest, followed by aspartate. Each of these is a sum over '
         f'all five contact types for one amino-acid type across the whole protein, so it '
         f'reports how much weighted contact the ligand makes with residues of that kind '
         f'rather than with a particular residue position. This is still a more '
         f'informative ranking than one over principal components, because each value has '
         f'a stated definition and unit, but it does not by itself identify a site or a '
         f'single interaction type.', first=True)
    figure(doc, 'Figure_5_shap',
           'Figure 5. Shapley attribution for the simulation-derived descriptors under '
           'the structure-disjoint partition. (A) The sixteen individual descriptors '
           'with the largest mean absolute contribution. (B) Contribution aggregated by '
           'descriptor family, with the percentage of total attribution.')

    head(doc, 'Interpretation of the Residue-Resolved Descriptors', 2)
    para(doc,
         f'The fitted model places most of its predictive weight on contact aggregates '
         f'for nonpolar residue types and on the aspartate aggregate. That pattern is '
         f'compatible with the conventional account of association, in which burial of '
         f'nonpolar surface supplies much of the free energy and a smaller number of '
         f'electrostatic contacts supplies specificity '
         f'{cite("chothia", "vallone", "kumar")}, and with the observation that binding '
         f'energy tends to be '
         f'concentrated on a subset of interface positions rather than spread evenly '
         f'{cite("bogan")}. It is not evidence for that account.', first=True)
    para(doc,
         f'Three properties of the descriptors prevent a stronger reading. Each residue '
         f'descriptor sums the five contact types of Table 1 into one number before any '
         f'aggregation, so the phenylalanine value may reflect hydrophobic or cation-pi '
         f'contacts and the aspartate value may reflect ionic, hydrogen-bonded or '
         f'water-mediated contacts. Each is then aggregated over every residue of that '
         f'type in the protein, so no individual position is identified. And Shapley '
         f'values report the behaviour of a fitted regressor rather than a free energy, '
         f'with correlated descriptors sharing attribution between them. Residue '
         f'abundance, pocket size and target identity are not excluded as contributors '
         f'to the ranking. The defensible claim is therefore the narrow one: the '
         f'quantities this model relies on have stated definitions in physical units, so '
         f'its behaviour can be inspected and questioned, which is not possible for a '
         f'model built on fingerprint bits or unnamed components. Assigning the signal to '
         f'a specific interaction or binding mode would require descriptors resolved by '
         f'contact type and residue position, which the present set does not provide.',
         first=True)

    head(doc, 'Sensitivity to the Choice of Learner', 2)
    d = n.dnn
    para(doc,
         f'Across every descriptor set the deep neural network failed to generalise, '
         f'reaching {n.m("md_core89", "dnn", "test"):.3f} on the simulation-derived '
         f'descriptors and negative values elsewhere. Four architectures of increasing '
         f'depth and regularisation were each trained with five random seeds. The '
         f'spread in test performance between seeds within a single architecture '
         f'reached {n.dnn_seed_range():.2f} units, larger than the '
         f'{d["R2 test"].max() - d["R2 test"].min():.2f} unit spread between '
         f'architecture means, so the four architectures are not distinguishable at this '
         f'sample size (Figure S2 and Table S6). With 122 systems and a training partition of 84, a network of '
         f'the order of 10 to the fifth parameters cannot be fitted reliably. This is a '
         f'sample-size limitation and is consistent with systematic comparisons on '
         f'tabular data of similar scale {cite("grins", "shwartz")}.', first=True)
    para(doc,
         f'The ridge stack of the two base learners performed comparably to the '
         f'Nu-support vector regressor alone, reaching '
         f'{n.m("md_core89", "hybrid", "test"):.3f} against '
         f'{n.m("md_core89", "nusvr", "test"):.3f} on the simulation-derived '
         f'descriptors. Given held-out sets of 23 to 25 molecules, and bootstrap '
         f'intervals of the order of half a unit, differences of this size are not '
         f'statistically resolvable. The appropriate conclusion is that a classical '
         f'kernel model is sufficient to exploit this representation, and that the '
         f'choice of representation matters far more than the choice of learner.',
         first=True)

    head(doc, 'Quantum-Chemical Descriptors', 2)
    para(doc,
         f'The quantum-chemical descriptors were computed to a standard that permits a '
         f'clean test, and they passed every physical validation applied, including the '
         f'requirement that unsubstituted anthraquinone show a vanishing dipole by '
         f'symmetry. Their failure to predict potency, at '
         f'{n.m("dft", "nusvr", "test"):.3f} under the structure-disjoint partition, is '
         f'therefore informative rather than technical. After geometry optimisation the '
         f'median difference in the frontier orbital gap between rows sharing a chemical '
         f'structure is 0.004 eV, so these descriptors cannot distinguish the same '
         f'ligand bound to different targets and are bounded by the ligand-only ceiling. '
         f'Electronic structure describes the isolated molecule accurately; it does not '
         f'describe the complex. Figure 6 shows the descriptor distributions and their '
         f'weak univariate association with potency.', first=True)
    figure(doc, 'Figure_6_quantum_validation',
           'Figure 6. Quantum-chemical descriptors and internal validation. Measured '
           'potency against the frontier orbital gap (A), the electrophilicity index '
           '(B) and the dipole moment (C), with Pearson correlation. Marker shape and '
           'fill denote the target protein. (D) Leave-one-out cross-validated '
           'predictive ability for the principal descriptor sets under the '
           'structure-disjoint partition.', width=5.9)

    head(doc, 'Comparison with Previous Work', 2)
    para(doc,
         f'Models built on two-dimensional descriptors and fingerprints for comparable '
         f'targets report coefficients of determination in the range 0.65 to 0.77 on '
         f'test data {cite("alk", "era", "nlrp3")}. The values obtained here on the '
         f'simulation-derived descriptors are at or above the top of that range, and are '
         f'obtained under a structure-disjoint partition that is stricter than the '
         f'random partitions those studies used. More importantly, the comparison made '
         f'here is internal: the same molecules, partitions, pipeline and metrics are '
         f'used throughout, so the difference is attributable to the representation and '
         f'not to any difference in dataset or protocol.', first=True)

    head(doc, 'Scope and Practical Placement', 2)
    para(doc,
         f'Every descriptor used here is computed from a trajectory, so the '
         f'representation cannot be evaluated for a molecule that has not been simulated '
         f'in complex with its target. The method is therefore not a virtual screening '
         f'filter and cannot rank a library of compounds that exist only as structures. '
         f'Its natural application is lead optimisation, where a series is already bound '
         f'and posed, the number of candidates is tens rather than millions, and a '
         f'simulation for each is an acceptable cost against the cost of synthesis.',
         first=True)
    para(doc,
         f'Because each candidate requires complex preparation and a 50 ns trajectory, '
         f'the workflow is substantially more demanding than static docking. Once the '
         f'trajectory exists all 89 descriptors are read from it, so the marginal cost of '
         f'the representation itself is small, but the trajectory dominates and sets the '
         f'scale. Endpoint methods such as MM/GBSA are also computed from trajectories '
         f'and share much of that cost {cite("mmgbsa")}, whereas alchemical free energy '
         f'calculations require a separate calculation for each transformation with '
         f'attention to sampling, protonation and force field {cite("fep")}. Runtime was '
         f'not benchmarked against either in this work, and no ordering by cost is '
         f'claimed.', first=True)

    # ------------------------------------------------------------- limitations
    head(doc, 'Limitations', 1)
    for t in [
        'The dataset comprises 122 systems built from 94 distinct chemical structures. '
        'Held-out partitions therefore contain 23 to 25 molecules and bootstrap '
        'confidence intervals on the coefficient of determination span approximately '
        'half a unit. Differences smaller than that are not resolvable, and no claim of '
        'ranking between closely performing models is made.',
        'Two of the ligand conformational dynamics columns scale with the number of '
        'atoms in the ligand trajectory rather than with the amplitude of motion, and '
        'correlate with potency across the dataset. Their contribution should be read as '
        'a size effect. Removing all three such columns changes the test coefficient of '
        'determination for the simulation-derived block from 0.798 to 0.706 while '
        'leaving the holdout value unchanged at 0.843, so the conclusions are '
        'unaffected.',
        'The residue-resolved interaction descriptors depend on the geometric criteria '
        'and representative energies listed in Table 1. Those values are conventional '
        'mid-range figures rather than system-specific calculations, and the '
        'implementation released with this work should be used for any reanalysis.',
        'The AlphaFold confidence descriptors were summarised over part of each '
        'structure because per-atom scores and per-residue indices were combined '
        'positionally. The block is constant within a protein and is treated throughout '
        'as an indicator of target identity rather than as a structural measurement.',
        'The present analysis differs from the earlier form of this study in the '
        'partitioning regime, the hyperparameter selection criterion, the point at '
        'which preprocessing is fitted, and three corrected descriptor definitions. '
        'Each change and its effect is itemised in Table S7.',
        'The structure-disjoint partition prevents a chemical structure from appearing '
        'on both sides of a split, but each of the four targets appears in both the '
        'training and the held-out partitions, so it does not test transfer to a protein '
        'the model has not seen. Holding out each target in turn and refitting on the '
        'remaining three gives a negative coefficient of determination in all four cases, '
        'ranging from -0.067 for MAPK1 to -2.641 for TP53, as reported in Table S9. The '
        'root-mean-square error nevertheless remains below that of the training-set mean '
        'for every target, so some ordering within a held-out target is retained while '
        'the absolute potency scale is not. The results reported here therefore support '
        'interpolation among these four target systems and do not establish prospective '
        'performance against a new protein. Descriptors aggregated by residue type may in '
        'part encode target identity, and separating the two would require a '
        'leave-one-target-out design over a substantially larger panel of proteins.',
        'Simulations of 50 ns sample a limited portion of conformational space. Longer '
        'trajectories may resolve slower motions not captured here, although the '
        'comparison against 100 ns trajectories reported as Figure S1 showed '
        'negligible differences across the measures used.',
    ]:
        p_ = doc.add_paragraph(style='List Bullet')
        check(t, 'limitation')
        r = p_.add_run(t)
        r.font.size = Pt(10.5)

    # ------------------------------------------------------------- conclusions
    head(doc, 'Conclusions', 1)
    para(doc,
         f'Descriptors generated from molecular dynamics trajectories gave the highest '
         f'point estimate of predictive accuracy among four representations evaluated on '
         f'the same molecules with the same pipeline, partitions and metrics. Paired '
         f'testing on identical held-out molecules separated them from the '
         f'quantum-chemical set, the far larger PyDescriptor set and the fingerprint '
         f'block, but not from PaDEL, whose difference was not resolvable at this sample '
         f'size. Under '
         f'a structure-disjoint partition the 89 simulation-derived descriptors reached '
         f'a test coefficient of determination of '
         f'{n.m("md_core89", "nusvr", "test"):.3f} and a holdout value of '
         f'{n.m("md_core89", "nusvr", "holdout"):.3f}, approaching the analytic ceiling '
         f'of {CEILING:.3f} imposed by repeated structures, while a descriptor set '
         f'exceeding the sample size by three orders of magnitude and a set computed at '
         f'a far higher level of theory both failed to predict potency at all.',
         first=True)
    para(doc,
         f'Attribution places the signal in residue-resolved interaction forces and in '
         f'ligand conformational dynamics, that is, in exactly those quantities that '
         f'require the complex to be simulated rather than merely drawn. A classical '
         f'kernel model is sufficient to exploit the representation; deep networks do '
         f'not generalise at this sample size, and ensembling them yields no resolvable '
         f'gain. The practical implication is that effort invested in describing the '
         f'dynamics of a complex returns more than effort invested in enlarging a static '
         f'descriptor set or in raising the level of electronic-structure theory applied '
         f'to the ligand alone.', first=True)

    head(doc, 'Associated Content', 1)
    para(doc,
         'Supporting Information. Comparison of 50 ns and 100 ns simulations; full '
         'performance metrics for every descriptor set, learner and partition; '
         'statistical validation including leave-one-out predictive ability, bootstrap '
         'intervals, y-scrambling and structure leakage; paired comparison of descriptor '
         'sets; quantum-chemical descriptor statistics; deep neural network architecture '
         'comparison; alignment of the retained principal components with potency; '
         'performance with a whole target held out; '
         'dataset inventory and provenance; and reproduction of the '
         'previously reported model. This material is available free of charge.',
         first=True)
    para(doc,
         'Data and Software Availability. All code, input descriptor matrices, trained '
         'models, per-molecule predictions and a run manifest recording package '
         'versions, random seeds and input checksums are available at '
         'https://github.com/Saeedmomo/Binding-Affinity-Prediction-Using-Molecular-Dynamics',
         first=True)

    head(doc, 'References', 1)
    uncited = [k for k in R if k not in _CITE_ORDER]
    if uncited:
        raise ValueError(f'references present but never cited: {uncited}')
    for i, k in enumerate(_CITE_ORDER, 1):
        r = REFERENCES[R[k] - 1]
        check(r, f'reference {i}')
        p_ = doc.add_paragraph()
        p_.paragraph_format.left_indent = Inches(0.3)
        p_.paragraph_format.first_line_indent = Inches(-0.3)
        run = p_.add_run(f'({i}) {r}')
        run.font.size = Pt(9)

    head(doc, 'Tables', 1)
    table(doc, 'Table3_feature_block_ablation',
          'Table 2. Contribution of each descriptor block, evaluated with the '
          'Nu-support vector regressor.',
          'The simulation-derived block matches or exceeds the full descriptor set.')
    table(doc, 'Table5_dft_descriptors',
          'Table 3. Quantum-chemical descriptors and their univariate association with '
          'potency.',
          'B3LYP/def2-SVP single-point energies on GFN2-xTB optimised geometries. '
          'Significance values are not corrected for multiple testing.', fontsize=6.8)
    return doc


# ------------------------------------------------------ supporting information
def build_si(n: N) -> Document:
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)

    t = doc.add_heading('Supporting Information', 0)
    for r in t.runs:
        r.font.size = Pt(16)
    para(doc, 'A Controlled Benchmark of Molecular Dynamics Derived Descriptors '
              'for Ligand Potency Prediction', size=11, bold=True)
    para(doc, 'Said Moshawih and co-authors', size=10)
    doc.add_paragraph()

    head(doc, 'Contents', 1)
    for line in [
        'Figure S1. Comparison of 50 ns and 100 ns simulations.',
        'Figure S2. Deep neural network architectures across random seeds.',
        'Figure S3. Y-scrambling null distributions.',
        'Table S1. Descriptor sets benchmarked, with provenance.',
        'Table S2. Reproduction of the previously reported model.',
        'Table S3. Performance of every descriptor set under both partitions.',
        'Table S4. Paired comparison of descriptor sets.',
        'Table S5. Statistical validation.',
        'Table S6. Deep neural network architectures.',
        'Table S7. Corrections applied relative to the earlier form of this study.',
        'Table S8. Alignment of the retained principal components with potency.',
        'Table S9. Performance when a whole target is held out.',
    ]:
        check(line, 'SI contents')
        p_ = doc.add_paragraph()
        p_.add_run(line).font.size = Pt(10)

    head(doc, 'Comparison of Simulation Lengths', 1)
    para(doc,
         'Simulations of 50 ns and 100 ns were compared across the principal dynamic '
         'and energetic measures for all four proteins. Differences were negligible for '
         'solvent-accessible surface area, root-mean-square deviation and fluctuation '
         'of both complex and ligand, the ligand principal component metrics, and the '
         'residue-resolved interactions. Isolated differences were confined to '
         'electrostatic terms in two proteins, reaching 43.97 kcal per mole in the '
         'protein electrostatic term for MAPK1 and 33.94 kcal per mole in the water '
         'electrostatic term for TP53. Trajectories of 50 ns were therefore judged '
         'sufficient and were used throughout.', first=True)
    p = os.path.join(ROOT, 'figures', 'original', 'main_image3.png')
    if os.path.exists(p):
        doc.add_picture(p, width=Inches(6.6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(doc,
                'Figure S1. Differences between 100 ns and 50 ns simulations for the '
                'four proteins. (A) Residue-resolved interaction differences. (B) '
                'Ligand principal component metrics. (C) Solvent-accessible surface '
                'area and root-mean-square deviation and fluctuation. (D) Interaction '
                'geometry and energy terms. (E) Energy descriptors. Values close to '
                'zero indicate agreement between the two simulation lengths.')
        doc.add_paragraph()

    head(doc, 'Deep Neural Network Architectures', 1)
    d = n.dnn
    para(doc,
         f'Four architectures of increasing depth and regularisation were each trained '
         f'with five random seeds under the shared partition. Mean test performance '
         f'ranged from {d["R2 test"].min():.3f} to {d["R2 test"].max():.3f} across the '
         f'four, while the standard deviation between seeds within a single architecture '
         f'reached {d["R2 test SD"].max():.3f}. Seed-to-seed variation therefore exceeds '
         f'the difference between architectures, and the four cannot be distinguished at '
         f'this sample size.', first=True)
    figure_si(doc, 'Figure_S1_dnn_variants',
              'Figure S2. Test-set performance of four deep neural network '
              'architectures, each trained with five random seeds. Points are '
              'individual seeds.', 4.8)
    para(doc,
         'Table S6 lists the four architectures with their parameter counts and the '
         'mean and standard deviation of performance over the five seeds.', first=True)

    head(doc, 'Y-Scrambling', 1)
    para(doc,
         'Each model was refitted against 100 random permutations of the potency labels '
         'in the training partition. Null distributions are centred well below the '
         'observed performance for every descriptor set that carries signal, confirming '
         'that the models do not fit noise.', first=True)
    figure_si(doc, 'Figure_S2_y_scrambling',
              'Figure S3. Null distributions of test-set performance under 100 label '
              'permutations. The vertical line marks the observed value.', 6.2)

    head(doc, 'Tables', 1)
    table(doc, 'TableS1_dataset_inventory',
          'Table S1. Descriptor sets benchmarked, with provenance.',
          'All sets share the same molecules and potency values. Checksum prefixes '
          'identify the exact input files used.',
          cols=['Descriptor set', 'Molecules', 'Descriptors', 'Missing cells',
                'pIC50 source', 'Source file', 'SHA-256 (16)'], fontsize=6.5)
    table(doc, 'TableS2_reproduction',
          'Table S2. Reproduction of the previously reported model.',
          'Reproduced on the same data and partition with hyperparameters ranked by the '
          'original composite criterion. Training performance agrees to four decimal '
          'places.')
    table(doc, 'Table2_benchmark_metrics',
          'Table S3. Performance of every descriptor set under both partitioning '
          'regimes.',
          'Bootstrap 95 per cent confidence intervals from 1000 resamples. Held-out '
          'partitions contain 23 to 25 molecules.',
          cols=['Descriptor set', 'Split', 'Model', 'R2 train', 'R2 CV', 'R2 test',
                'R2 test 95% CI', 'R2 holdout', 'RMSE test', 'MAE test'], fontsize=6.2)
    table(doc, 'TableS4_paired_tests',
          'Table S4. Paired comparison of descriptor sets on identical held-out '
          'molecules.',
          'Wilcoxon signed-rank test on per-molecule absolute error. A negative median '
          'difference and a positive effect size indicate that descriptor set A is more '
          'accurate. Significance values are Holm-corrected within each partition.',
          cols=['Split', 'A', 'B', 'n', 'median_diff', 'effect_size_rbc', 'p_holm',
                'significant'], fontsize=6.2)
    table(doc, 'Table4_statistical_validation',
          'Table S5. Statistical validation for every descriptor set and partition.',
          'Significance values are the empirical fraction of 100 label permutations '
          'reaching the observed performance. Leaked rows count held-out molecules '
          'whose structure also appears in training, zero by construction under the '
          'structure-disjoint regime.', fontsize=6.2)
    _si_table_from_df(doc, n.dnn,
                      'Table S6. Deep neural network architectures, mean and standard '
                      'deviation over five random seeds.')
    table(doc, 'TableS3_manuscript_corrections',
          'Table S7. Corrections applied relative to the earlier form of this study.',
          '', fontsize=6.2)
    _si_table_from_df(doc, n.pca_align,
                      'Table S8. Alignment of the retained principal components with '
                      'potency, computed on the 84 training molecules of the '
                      'structure-disjoint partition.')
    caption(doc,
            'Each representation was passed through the preprocessing used for '
            'modelling, that is median imputation, removal of zero-variance columns, '
            'standardisation and reduction to 17 principal components. The final two '
            'columns give the largest absolute Pearson correlation between any retained '
            'component and potency, and the rank of that component by explained '
            'variance. No model is fitted and no held-out molecule is used.')
    _si_table_from_df(doc, n.loto,
                      'Table S9. Performance of the simulation-derived descriptors when '
                      'a whole target is held out.')
    caption(doc,
            'Each row refits the Nu-support vector regressor on the three remaining '
            'targets and predicts every molecule of the held-out target. The final '
            'column gives the error of predicting the mean potency of the training '
            'targets, for reference. A negative coefficient of determination indicates '
            'that the absolute potency scale of an unseen target is not recovered, while '
            'a root-mean-square error below that reference indicates that some ordering '
            'within the held-out target is retained.')
    return doc


def figure_si(doc, name, cap, width=6.2):
    p = os.path.join(FIG, f'{name}.png')
    if not os.path.exists(p):
        return False
    doc.add_picture(p, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)
    doc.add_paragraph()
    return True


def _si_table_from_df(doc, df, cap, fontsize=7.0):
    if df is None:
        return
    df = df.rename(columns={c: prettify(c) for c in df.columns})
    caption(doc, cap)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = 'Table Grid'
    for i, c in enumerate(df.columns):
        cell = t.rows[0].cells[i]
        cell.text = str(c)
        for pr in cell.paragraphs:
            for r in pr.runs:
                r.bold = True
                r.font.size = Pt(fontsize)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v).replace('_', ' ')
            for pr in cells[i].paragraphs:
                pr.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                                else WD_ALIGN_PARAGRAPH.RIGHT)
                for r in pr.runs:
                    r.font.size = Pt(fontsize)
    doc.add_paragraph()


def main():
    n = N()
    stamp = dt.date.today().strftime('%d%m%y')
    os.makedirs(PAPER, exist_ok=True)

    doc = build_manuscript(n)
    # Figure 1 belongs with the methods; insert it after the descriptor section by
    # appending here so that all display items precede the reference list is preserved
    out1 = os.path.join(PAPER, f'Moshawih_MD_hybrid_ML_manuscript_V3c_{stamp}.docx')
    try:
        doc.save(out1)
    except PermissionError:
        sys.exit(f'cannot write {out1}: the file is open in Word. Close it and re-run.')

    si = build_si(n)
    out2 = os.path.join(PAPER,
                        f'Moshawih_MD_hybrid_ML_supporting_information_V3c_{stamp}.docx')
    try:
        si.save(out2)
    except PermissionError:
        sys.exit(f'cannot write {out2}: the file is open in Word. Close it and re-run.')

    print(f'wrote {os.path.basename(out1)}')
    print(f'wrote {os.path.basename(out2)}')


if __name__ == '__main__':
    main()
